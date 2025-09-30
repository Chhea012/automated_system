from flask import Blueprint, render_template, request, redirect, url_for, flash, send_file, jsonify
from flask_login import login_required, current_user
from app import db
from app.models.contract import Contract
from app.models.notification import Notification
from app.models.user import User
import uuid
from datetime import datetime
import pandas as pd
from io import BytesIO
import logging
from num2words import num2words
import re
from openpyxl import Workbook
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor
import zipfile
from docx.enum.text import WD_TAB_ALIGNMENT

# Configure logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

contracts_bp = Blueprint('contracts', __name__)

def sanitize_filename(name):
    return re.sub(r'[^\w\s.-]', ' ', name.replace(' ', ' ')).strip()

def generate_next_contract_number(last_contract_number, current_year):
    if not last_contract_number:
        return f"NGOF/{current_year}-001"
    try:
        match = re.match(r"NGOF/(\d{4})-(\d{3})", last_contract_number)
        if not match:
            logger.error(f"Invalid contract number format: {last_contract_number}")
            return f"NGOF/{current_year}-001"
        year, number = match.groups()
        if year == str(current_year):
            next_number = int(number) + 1
            return f"NGOF/{year}-{next_number:03d}"
        else:
            return f"NGOF/{current_year}-001"
    except Exception as e:
        logger.error(f"Error generating next contract number: {str(e)}")
        return f"NGOF/{current_year}-001"

def format_date(iso_date):
    try:
        if not iso_date or iso_date.lower() in ['n/a', '']:
            return ''
        if 'week' in iso_date.lower():
            return iso_date
        date = datetime.strptime(iso_date, '%Y-%m-%d')
        day = date.day
        month = date.strftime('%B')
        year = date.year

        # Determine suffix
        if 11 <= day % 100 <= 13:
            suffix = 'th'
        else:
            suffix = {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')

        # Map to Unicode superscripts
        superscripts = {
            "st": "ˢᵗ",
            "nd": "ⁿᵈ",
            "rd": "ʳᵈ",
            "th": "ᵗʰ"
        }

        return f"{day}{superscripts[suffix]} {month} {year}"
    except (ValueError, TypeError) as e:
        logger.warning(f"Error formatting date '{iso_date}': {str(e)}")
        return iso_date or ''

def number_to_words(num):
    try:
        if not num or num < 0:
            return "Zero US Dollars only"
        integer_part = int(num)
        decimal_part = round((num - integer_part) * 100)
        words = num2words(integer_part, lang='en').title()
        if decimal_part > 0:
            words += " and " + num2words(decimal_part, lang='en').title() + " Cents"
        return f"{words} US Dollars only"
    except Exception as e:
        logger.error(f"Error converting number to words: {str(e)}")
        return "N/A"

def normalize_to_list(field):
    if isinstance(field, list):
        return [str(item).strip() for item in field if str(item).strip()]
    elif isinstance(field, str):
        return [item.strip() for item in field.split('\n') if item.strip()]
    return []

def calculate_installment_payments(total_fee_usd, tax_percentage, percentage):
    try:
        gross_amount = (total_fee_usd * percentage) / 100
        tax_amount = gross_amount * (tax_percentage / 100)
        net_amount = gross_amount - tax_amount
        return gross_amount, tax_amount, net_amount
    except Exception as e:
        logger.error(f"Error calculating installment payments: {str(e)}")
        return 0.0, 0.0, 0.0

def calculate_payments(total_fee_usd, tax_percentage, payment_installments):
    try:
        total_gross = 0.0
        total_net = 0.0
        for installment in payment_installments:
            match = re.search(r'\((\d+\.?\d*)\%\)', installment['description'])
            if not match:
                logger.warning(f"Invalid percentage format in installment: {installment['description']}")
                continue
            percentage = float(match.group(1))
            gross_amount = (total_fee_usd * percentage) / 100
            net_amount = gross_amount * (1 - tax_percentage / 100)
            total_gross += gross_amount
            total_net += net_amount
        return total_gross, total_net
    except Exception as e:
        logger.error(f"Error calculating payments: {str(e)}")
        return 0.0, 0.0
#list of the contract
@contracts_bp.route('/')
@login_required
def index():
    try:
        # Mark notifications as read for Admins
        if current_user.has_role('admin'):
            Notification.query.filter_by(recipient_id=current_user.id, is_read=False).update({'is_read': True})
            db.session.commit()
            logger.info(f"Notifications marked as read for user {current_user.id}")

        page = request.args.get('page', 1, type=int)
        search_query = request.args.get('search', '', type=str)
        sort_order = request.args.get('sort', 'created_at_desc', type=str)
        entries_per_page = request.args.get('entries', 10, type=int)

        query = Contract.query.filter(Contract.deleted_at == None)
        if not current_user.has_role('admin'):
            query = query.filter(Contract.user_id == current_user.id)

        if search_query:
            query = query.filter(
                (Contract.project_title.ilike(f'%{search_query}%')) |
                (Contract.contract_number.ilike(f'%{search_query}%')) |
                (Contract.party_b_signature_name.ilike(f'%{search_query}%'))
            )

        if sort_order == 'contract_number_asc':
            query = query.order_by(Contract.contract_number.asc())
        elif sort_order == 'contract_number_desc':
            query = query.order_by(Contract.contract_number.desc())
        elif sort_order == 'start_date_asc':
            query = query.order_by(Contract.agreement_start_date.asc())
        elif sort_order == 'start_date_desc':
            query = query.order_by(Contract.agreement_start_date.desc())
        elif sort_order == 'total_fee_asc':
            query = query.order_by(Contract.total_fee_usd.asc())
        elif sort_order == 'total_fee_desc':
            query = query.order_by(Contract.total_fee_usd.desc())
        else:
            query = query.order_by(Contract.created_at.desc())

        pagination = query.paginate(page=page, per_page=entries_per_page, error_out=False)
        contracts = [contract.to_dict() for contract in pagination.items]

        for contract in contracts:
            contract['agreement_start_date_display'] = format_date(contract.get('agreement_start_date'))
            contract['agreement_end_date_display'] = format_date(contract.get('agreement_end_date'))
            contract['total_fee_usd'] = f"{contract.get('total_fee_usd', 0.0):.2f}"
            if 'custom_article_sentences' not in contract or contract['custom_article_sentences'] is None:
                contract['custom_article_sentences'] = []

        total_contracts = query.count()
        total_contracts_global = Contract.query.filter(Contract.deleted_at == None).count()
        last_contract = Contract.query.filter(Contract.deleted_at == None).order_by(Contract.contract_number.desc()).first()
        last_contract_number = last_contract.contract_number if last_contract else None

        return render_template(
            'contracts/index.html',
            contracts=contracts,
            pagination=pagination,
            search_query=search_query,
            sort_order=sort_order,
            entries_per_page=entries_per_page,
            total_contracts=total_contracts,
            total_contracts_global=total_contracts_global,
            last_contract_number=last_contract_number,
            is_admin=current_user.has_role('admin')
        )
    except Exception as e:
        logger.error(f"Error in index route: {str(e)}")
        flash("An error occurred while loading contracts.", 'danger')
        return render_template(
            'contracts/index.html',
            contracts=[],
            pagination=None,
            search_query='',
            sort_order='created_at_desc',
            entries_per_page=10,
            total_contracts=0,
            total_contracts_global=0,
            last_contract_number=None,
            is_admin=current_user.has_role('admin')
        )    

@contracts_bp.route('/create', methods=['GET', 'POST'])
@login_required
def create():
    current_year = datetime.now().year
    last_contract = Contract.query.filter(Contract.deleted_at == None).order_by(Contract.contract_number.desc()).first()
    last_contract_number = last_contract.contract_number if last_contract else None
    default_contract_number = generate_next_contract_number(last_contract_number, current_year)

    # Fetch unique Party A data from previous contracts
    previous_contracts = Contract.query.filter(Contract.deleted_at == None).all()
    party_a_data = {}
    for contract in previous_contracts:
        for person in contract.party_a_info or []:
            if isinstance(person, dict) and person.get('name'):
                name = person['name'].strip()
                normalized_name = name.lower()
                if name and normalized_name not in party_a_data:
                    party_a_data[normalized_name] = {
                        'name': name,
                        'position': person.get('position', '').strip(),
                        'address': person.get('address', '').strip()
                    }

    # Fetch unique Party B data
    party_b_data = {}
    for contract in previous_contracts:
        name = contract.party_b_signature_name.strip()
        if name and name.lower() not in party_b_data:
            party_b_data[name.lower()] = {
                'original_name': name,
                'position': contract.party_b_position or '',
                'phone': contract.party_b_phone or '',
                'email': contract.party_b_email or '',
                'address': contract.party_b_address or ''
            }

    # Fetch unique focal person data
    focal_person_data = {}
    for contract in previous_contracts:
        focal_persons = contract.focal_person_info or [] 
        for person in focal_persons:
            if isinstance(person, dict) and person.get('name'):
                name = person['name'].strip()
                normalized_name = name.lower()
                if name and normalized_name not in focal_person_data:
                    focal_person_data[normalized_name] = {
                        'name': name,
                        'position': person.get('position', '').strip(),
                        'phone': person.get('phone', '').strip(),
                        'email': person.get('email', '').strip()
                    }

    form_data = {}
    if request.method == 'POST':
        try:
            # Collect simple fields
            party_b_select = request.form.get('party_b_select', '').strip()
            party_b_name = request.form.get('party_b_signature_name', '').strip() if party_b_select == 'new' else party_b_select
            party_a_signer = request.form.get('party_a_signer', '').strip()
            deduct_tax_code = request.form.get('deduct_tax_code', '').strip()
            vat_organization_name = request.form.get('vat_organization_name', '').strip()  # New field

            form_data = {
                'project_title': request.form.get('project_title', '').strip(),
                'contract_number': request.form.get('contract_number', '').strip(),
                'output_description': request.form.get('output_description', '').strip(),
                'tax_percentage': float(request.form.get('tax_percentage', '15.0').strip() or 15.0),
                'deduct_tax_code': deduct_tax_code,
                'vat_organization_name': vat_organization_name,  # New in form_data
                'organization_name': request.form.get('organization_name', '').strip(),
                'party_b_signature_name': party_b_name,
                'party_b_position': request.form.get('party_b_position', '').strip(),
                'party_b_phone': request.form.get('party_b_phone', '').strip(),
                'party_b_email': request.form.get('party_b_email', '').strip(),
                'party_b_address': request.form.get('party_b_address', '').strip(),
                'agreement_start_date': request.form.get('agreement_start_date', '').strip(),
                'agreement_end_date': request.form.get('agreement_end_date', '').strip(),
                'total_fee_usd': float(request.form.get('total_fee_usd', '0.0').strip() or 0.0),
                'total_fee_words': request.form.get('total_fee_words', '').strip(),
                'workshop_description': request.form.get('workshop_description', '').strip(),
                'title': request.form.get('title', '').strip(),
                'party_b_full_name_with_title': party_b_name,
                'party_b_signature_name_confirm': request.form.get('party_b_signature_name_confirm', '').strip(),
                'party_b_select': party_b_select,
                'party_a_signer': party_a_signer
            }

            # Process Party A info (multiple entries)
            party_a_info = [
                {
                    'name': name.strip(),
                    'position': pos.strip(),
                    'address': addr.strip()
                }
                for name, pos, addr in zip(
                    request.form.getlist('party_a_name[]'),
                    request.form.getlist('party_a_position[]'),
                    request.form.getlist('party_a_address[]')
                )
                if name.strip() and pos.strip() and addr.strip()
            ]
            if not party_a_info:
                flash('At least one Party A representative is required.', 'danger')
                form_data['payment_installments'] = []
                form_data['focal_person_info'] = []
                form_data['articles'] = []
                form_data['party_a_info'] = [{'name': '', 'position': '', 'address': ''}]
                return render_template('contracts/create.html', form_data=form_data, default_contract_number=default_contract_number, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)

            form_data['party_a_info'] = party_a_info

            # Validate Party A signer
            if not party_a_signer or party_a_signer not in [p['name'] for p in party_a_info]:
                flash('Please select a valid Party A signer from the list.', 'danger')
                form_data['payment_installments'] = []
                form_data['focal_person_info'] = []
                form_data['articles'] = []
                return render_template('contracts/create.html', form_data=form_data, default_contract_number=default_contract_number, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)

            # Validate Party B name
            if not party_b_name or not re.match(r'^[a-zA-Z\s\.]+$', party_b_name):
                flash('Party B signature name is required and must contain only letters, spaces, and periods.', 'danger')
                form_data['payment_installments'] = []
                form_data['focal_person_info'] = []
                form_data['articles'] = []
                return render_template('contracts/create.html', form_data=form_data, default_contract_number=default_contract_number, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)

            # Validate deduct_tax_code and vat_organization_name when tax_percentage is 0
            if form_data['tax_percentage'] == 0:
                if not deduct_tax_code:
                    flash('VAT TIN is required when tax percentage is 0%.', 'danger')
                    form_data['payment_installments'] = []
                    form_data['focal_person_info'] = []
                    form_data['articles'] = []
                    return render_template('contracts/create.html', form_data=form_data, default_contract_number=default_contract_number, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)
                if not re.match(r'^[A-Z0-9\-]+$', deduct_tax_code):
                    flash('VAT TIN must contain only uppercase letters, numbers, and hyphens.', 'danger')
                    form_data['payment_installments'] = []
                    form_data['focal_person_info'] = []
                    form_data['articles'] = []
                    return render_template('contracts/create.html', form_data=form_data, default_contract_number=default_contract_number, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)
                if len(deduct_tax_code) > 50:
                    flash('VAT TIN must not exceed 50 characters.', 'danger')
                    form_data['payment_installments'] = []
                    form_data['focal_person_info'] = []
                    form_data['articles'] = []
                    return render_template('contracts/create.html', form_data=form_data, default_contract_number=default_contract_number, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)
                if not vat_organization_name:
                    flash('Name of Organization is required when tax percentage is 0%.', 'danger')
                    form_data['payment_installments'] = []
                    form_data['focal_person_info'] = []
                    form_data['articles'] = []
                    return render_template('contracts/create.html', form_data=form_data, default_contract_number=default_contract_number, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)
                if len(vat_organization_name) > 255:
                    flash('Name of Organization must not exceed 255 characters.', 'danger')
                    form_data['payment_installments'] = []
                    form_data['focal_person_info'] = []
                    form_data['articles'] = []
                    return render_template('contracts/create.html', form_data=form_data, default_contract_number=default_contract_number, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)

            # Process custom articles
            articles_raw = [
                {'article_number': num.strip(), 'custom_sentence': sent.strip()}
                for num, sent in zip(request.form.getlist('articleNumber[]'), request.form.getlist('customSentence[]'))
                if sent.strip()
            ]
            form_data['articles'] = articles_raw
            form_data['custom_article_sentences'] = {str(article['article_number']): article['custom_sentence'] for article in articles_raw}

            # Process payment installments
            payment_installments_raw = [
                {
                    'description': desc.strip(),
                    'deliverables': deliv.strip(),
                    'dueDate': due.strip()
                }
                for desc, deliv, due in zip(
                    request.form.getlist('paymentInstallmentDesc[]'),
                    request.form.getlist('paymentInstallmentDeliverables[]'),
                    request.form.getlist('paymentInstallmentDueDate[]')
                )
                if desc.strip() and deliv.strip() and due.strip()
            ]
            if not payment_installments_raw:
                flash('At least one payment installment is required.', 'danger')
                form_data['payment_installments'] = []
                form_data['focal_person_info'] = []
                return render_template('contracts/create.html', form_data=form_data, default_contract_number=default_contract_number, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)

            form_data['payment_installments'] = payment_installments_raw
            deliverables = '; '.join([inst['deliverables'] for inst in payment_installments_raw])
            form_data['deliverables'] = deliverables

            # Process focal persons
            focal_person_raw = [
                {
                    'name': name.strip(),
                    'position': pos.strip(),
                    'phone': phone.strip(),
                    'email': email.strip()
                }
                for name, pos, phone, email in zip(
                    request.form.getlist('focal_person_name[]'),
                    request.form.getlist('focal_person_position[]'),
                    request.form.getlist('focal_person_phone[]'),
                    request.form.getlist('focal_person_email[]')
                )
                if name.strip() and pos.strip() and phone.strip() and email.strip()
            ]
            if not focal_person_raw:
                flash('At least one focal person is required.', 'danger')
                form_data['focal_person_info'] = []
                return render_template('contracts/create.html', form_data=form_data, default_contract_number=default_contract_number, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)

            form_data['focal_person_info'] = focal_person_raw

            # Calculate payments
            total_fee_usd = form_data['total_fee_usd']
            tax_percentage = form_data['tax_percentage']
            gross_amount_usd = total_fee_usd
            total_gross, total_net = calculate_payments(total_fee_usd, tax_percentage, payment_installments_raw)
            form_data['payment_gross'] = f"${total_gross:.2f} USD"
            form_data['payment_net'] = f"${total_net:.2f} USD"
            form_data['gross_amount_usd'] = gross_amount_usd

            # Validate required fields
            required_fields = [
                ('project_title', 'Project title is required.'),
                ('contract_number', 'Contract number is required.'),
                ('output_description', 'Output description is required.'),
                ('organization_name', 'Organization name is required.'),
                ('party_b_signature_name', 'Party B signature name is required.'),
                ('agreement_start_date', 'Agreement start date is required.'),
                ('agreement_end_date', 'Agreement end date is required.'),
                ('total_fee_usd', 'Total fee USD is required.')
            ]
            for field, message in required_fields:
                if not form_data[field]:
                    flash(message, 'danger')
                    return render_template('contracts/create.html', form_data=form_data, default_contract_number=default_contract_number, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)

            # Validate Party B confirm match
            if form_data['party_b_signature_name'] != form_data['party_b_signature_name_confirm']:
                flash('Party B signature name confirmation does not match.', 'danger')
                return render_template('contracts/create.html', form_data=form_data, default_contract_number=default_contract_number, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)

            # Validate contract number format
            if not re.match(r"NGOF/\d{4}-\d{3}", form_data['contract_number']):
                flash('Contract number must follow the format NGOF/YYYY-NNN (e.g., NGOF/2025-005).', 'danger')
                return render_template('contracts/create.html', form_data=form_data, default_contract_number=default_contract_number, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)

            # Check for duplicate contract number
            if Contract.query.filter(Contract.contract_number == form_data['contract_number'], Contract.deleted_at == None).first():
                flash('Contract number already exists.', 'danger')
                return render_template('contracts/create.html', form_data=form_data, default_contract_number=default_contract_number, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)

            # Validate dates
            start_date = form_data['agreement_start_date']
            end_date = form_data['agreement_end_date']
            if start_date and end_date:
                try:
                    if datetime.strptime(end_date, '%Y-%m-%d') < datetime.strptime(start_date, '%Y-%m-%d'):
                        flash('Agreement end date must be after start date.', 'danger')
                        return render_template('contracts/create.html', form_data=form_data, default_contract_number=default_contract_number, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)
                except ValueError:
                    flash('Invalid date format for agreement start or end date.', 'danger')
                    return render_template('contracts/create.html', form_data=form_data, default_contract_number=default_contract_number, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)

            # Validate total_fee_usd
            if total_fee_usd < 0:
                flash('Total fee USD cannot be negative.', 'danger')
                return render_template('contracts/create.html', form_data=form_data, default_contract_number=default_contract_number, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)

            # Validate tax_percentage
            if tax_percentage not in [0, 5, 10, 15, 20]:
                flash('Tax percentage must be one of 0, 5, 10, 15, or 20.', 'danger')
                return render_template('contracts/create.html', form_data=form_data, default_contract_number=default_contract_number, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)

            # Validate payment installment percentages
            total_percentage = 0.0
            for installment in form_data['payment_installments']:
                match = re.search(r'\((\d+\.?\d*)\%\)', installment['description'])
                if not match:
                    flash(f"Invalid installment description format: {installment['description']}. Must include percentage like (50%).", 'danger')
                    return render_template('contracts/create.html', form_data=form_data, default_contract_number=default_contract_number, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)
                try:
                    percentage = float(match.group(1))
                    total_percentage += percentage
                except ValueError:
                    flash(f"Invalid percentage in installment description: {installment['description']}.", 'danger')
                    return render_template('contracts/create.html', form_data=form_data, default_contract_number=default_contract_number, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)
                try:
                    datetime.strptime(installment['dueDate'], '%Y-%m-%d')
                except ValueError:
                    flash(f"Invalid due date for installment: {installment['dueDate']}.", 'danger')
                    return render_template('contracts/create.html', form_data=form_data, default_contract_number=default_contract_number, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)

            if abs(total_percentage - 100.0) > 0.01:
                flash('Total percentage of payment installments must equal 100%.', 'danger')
                return render_template('contracts/create.html', form_data=form_data, default_contract_number=default_contract_number, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)

            # Validate focal person info
            for person in form_data['focal_person_info']:
                if not re.match(r'^[a-zA-Z\s\.]+$', person['name']):
                    flash(f"Invalid focal person name: {person['name']}. Only letters, spaces, and periods are allowed.", 'danger')
                    return render_template('contracts/create.html', form_data=form_data, default_contract_number=default_contract_number, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)
                if not re.match(r'^[a-zA-Z\s]+$', person['position']):
                    flash(f"Invalid focal person position: {person['position']}. Only letters and spaces are allowed.", 'danger')
                    return render_template('contracts/create.html', form_data=form_data, default_contract_number=default_contract_number, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)
                if not re.match(r'^\+?\d{1,4}([-.\s]?\d{1,4}){2,3}$', person['phone']):
                    flash(f"Invalid focal person phone: {person['phone']}. Use format like 012 845 091, +855 12 845 091, or +85512845091.", 'danger')
                    return render_template('contracts/create.html', form_data=form_data, default_contract_number=default_contract_number, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)
                if not re.match(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', person['email']):
                    flash(f"Invalid focal person email: {person['email']}.", 'danger')
                    return render_template('contracts/create.html', form_data=form_data, default_contract_number=default_contract_number, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)

            # Validate Party A info
            for person in form_data['party_a_info']:
                if not re.match(r'^[a-zA-Z\s\.]+$', person['name']):
                    flash(f"Invalid Party A name: {person['name']}. Only letters, spaces, and periods are allowed.", 'danger')
                    return render_template('contracts/create.html', form_data=form_data, default_contract_number=default_contract_number, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)
                if not re.match(r'^[a-zA-Z\s]+$', person['position']):
                    flash(f"Invalid Party A position: {person['position']}. Only letters and spaces are allowed.", 'danger')
                    return render_template('contracts/create.html', form_data=form_data, default_contract_number=default_contract_number, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)
                if not person['address']:
                    flash(f"Party A address is required.", 'danger')
                    return render_template('contracts/create.html', form_data=form_data, default_contract_number=default_contract_number, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)

            # Create new contract
            contract = Contract(
                id=str(uuid.uuid4()),
                user_id=current_user.id,
                project_title=form_data['project_title'],
                contract_number=form_data['contract_number'],
                organization_name=form_data['organization_name'],
                party_a_info=form_data['party_a_info'],
                party_b_full_name_with_title=form_data['party_b_full_name_with_title'],
                party_b_address=form_data['party_b_address'],
                party_b_phone=form_data['party_b_phone'],
                party_b_email=form_data['party_b_email'],
                registration_number='#304 សជណ',
                registration_date='07 March 2012',
                agreement_start_date=form_data['agreement_start_date'],
                agreement_end_date=form_data['agreement_end_date'],
                total_fee_usd=form_data['total_fee_usd'],
                gross_amount_usd=form_data['gross_amount_usd'],
                tax_percentage=form_data['tax_percentage'],
                deduct_tax_code=form_data['deduct_tax_code'] if form_data['tax_percentage'] == 0 else None,
                vat_organization_name=form_data['vat_organization_name'] if form_data['tax_percentage'] == 0 else None,  # New field
                payment_gross=form_data['payment_gross'],
                payment_net=form_data['payment_net'],
                workshop_description=form_data['workshop_description'],
                focal_person_info=form_data['focal_person_info'],
                party_a_signature_name=form_data['party_a_signer'],
                party_b_signature_name=form_data['party_b_signature_name'],
                party_b_position=form_data['party_b_position'],
                total_fee_words=form_data['total_fee_words'] or number_to_words(form_data['total_fee_usd']),
                title=form_data['title'],
                deliverables=form_data['deliverables'],
                output_description=form_data['output_description'],
                custom_article_sentences=form_data['custom_article_sentences'],
                payment_installments=form_data['payment_installments']
            )

            db.session.add(contract)
            db.session.commit()

            # Send notifications to all Admins (including creator)
            admins = User.query.filter(User.role.has(name='admin')).all()
            for admin in admins:
                notification = Notification(
                    creator_id=current_user.id,
                    recipient_id=admin.id,
                    title=f"New Contract Created: {contract.project_title}",
                    message=f"Contract {contract.contract_number} created by {current_user.username}",
                    related_contract_id=contract.id
                )
                db.session.add(notification)
            db.session.commit()
            logger.info(f"Notifications sent for contract {contract.contract_number} to {len(admins)} admins")

            flash('Contract created successfully!', 'success')
            return redirect(url_for('contracts.index'))
        except Exception as e:
            logger.error(f"Error creating contract: {str(e)}")
            flash("An error occurred while creating the contract.", 'danger')
            return render_template('contracts/create.html', form_data=form_data, default_contract_number=default_contract_number, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)

    # Initialize form_data for GET request
    form_data = {
        'party_a_info': [{'name': 'Mr. SOEUNG Saroeun', 'position': 'Executive Director', 'address': '#9-11, Street 476, Sangkat Tuol Tumpoung I, Phnom Penh, Cambodia'}],
        'focal_person_info': [{'name': '', 'position': '', 'phone': '', 'email': ''}],
        'payment_installments': [{'description': '', 'deliverables': '', 'dueDate': ''}],
        'articles': [],
        'custom_article_sentences': {},
        'party_a_signer': 'Mr. SOEUNG Saroeun',
        'deduct_tax_code': '',
        'vat_organization_name': ''  # New in init
    }
    return render_template('contracts/create.html', form_data=form_data, default_contract_number=default_contract_number, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)

@contracts_bp.route('/mark-read', methods=['POST'])
@login_required
def mark_notifications_read():
    try:
        if not current_user.has_role('admin'):
            return jsonify({'error': 'Unauthorized'}), 403
        Notification.query.filter_by(recipient_id=current_user.id, is_read=False).update({'is_read': True})
        db.session.commit()
        logger.info(f"Notifications marked as read via AJAX for user {current_user.id}")
        return jsonify({'success': True, 'unread_count': 0})
    except Exception as e:
        logger.error(f"Error marking notifications as read: {str(e)}")
        return jsonify({'error': 'Internal server error'}), 500    
#export docx file
@contracts_bp.route('/export_docx/<contract_id>')
@login_required
def export_docx(contract_id):
    try:
        contract = Contract.query.get_or_404(contract_id)
        # Allow admins to export any contract, non-admins only their own
        if not current_user.has_role('admin') and contract.user_id != current_user.id:
            flash("You are not authorized to export this contract.", 'danger')
            return redirect(url_for('contracts.index'))
        if contract.deleted_at is not None:
            flash("This contract has been deleted and cannot be exported.", 'danger')
            return redirect(url_for('contracts.index'))

        contract_data = contract.to_dict()
        if 'custom_article_sentences' not in contract_data or contract_data['custom_article_sentences'] is None:
            contract_data['custom_article_sentences'] = {}

        # Format dates
        contract_data['agreement_start_date_display'] = format_date(contract_data['agreement_start_date'])
        contract_data['agreement_end_date_display'] = format_date(contract_data['agreement_end_date'])

        # Get financial data as floats
        try:
            total_fee_usd = float(contract_data['total_fee_usd']) if contract_data['total_fee_usd'] else 0.0
            tax_percentage = float(contract_data.get('tax_percentage', 15.0))
            deduct_tax_code = contract_data.get('deduct_tax_code', '')
            vat_organization_name = contract_data.get('vat_organization_name', '')  # New field
        except (ValueError, TypeError) as e:
            logger.error(f"Error converting financial data for contract {contract_id}: {str(e)}")
            flash("An error occurred while exporting the contract.", 'danger')
            return redirect(url_for('contracts.index'))

        contract_data['total_fee_usd'] = total_fee_usd
        contract_data['gross_amount_usd'] = total_fee_usd
        contract_data['total_fee_words'] = contract_data.get('total_fee_words') or number_to_words(total_fee_usd)

        # Calculate total gross and net
        total_gross_amount, total_net_amount = calculate_payments(
            total_fee_usd, tax_percentage, contract_data.get('payment_installments', [])
        )
        contract_data['total_gross'] = f"USD{total_gross_amount:.2f}"
        contract_data['total_net'] = f"USD{total_net_amount:.2f}"

        # Process payment installments
        for installment in contract_data.get('payment_installments', []):
            installment['dueDate_display'] = format_date(installment.get('dueDate', ''))
            match = re.search(r'\((\d+\.?\d*)\%\)', installment['description'])
            percentage = float(match.group(1)) if match else 0.0
            gross, tax, net = calculate_installment_payments(total_fee_usd, tax_percentage, percentage)
            installment['gross_amount'] = gross
            installment['tax_amount'] = tax
            installment['net_amount'] = net

        # Create DOCX document
        doc = Document()

        # Set document margins and add footer to each section
        sections = doc.sections
        for i, section in enumerate(sections):
            if i == 0:  # Only modify the first section for letterhead space
                section.top_margin = Inches(1.2)  # Reduced to 1.2 inches for a more balanced letterhead space
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
                section.bottom_margin = Inches(1)
            else:
                section.top_margin = Inches(1)  # Default margin for subsequent pages
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
                section.bottom_margin = Inches(1)

            # Add footer with "Page X of Y"
            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Changed to right-aligned
            footer_para.paragraph_format.space_before = Pt(0)
            footer_para.paragraph_format.space_after = Pt(0)
            run = footer_para.add_run()
            run.font.name = 'Calibri'
            run.font.size = Pt(10)

            # Add "Page" text
            run.add_text('Page ')

            # Add page number field
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            run._r.append(fldChar1)
            instrText = OxmlElement('w:instrText')
            instrText.text = 'PAGE'
            run._r.append(instrText)
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            run._r.append(fldChar2)

            run.add_text(' of ')

            # Add total pages field
            fldChar3 = OxmlElement('w:fldChar')
            fldChar3.set(qn('w:fldCharType'), 'begin')
            run._r.append(fldChar3)
            instrText2 = OxmlElement('w:instrText')
            instrText2.text = 'NUMPAGES'
            run._r.append(instrText2)
            fldChar4 = OxmlElement('w:fldChar')
            fldChar4.set(qn('w:fldCharType'), 'end')
            run._r.append(fldChar4)

        # Set default font
        doc.styles['Normal'].font.name = 'Calibri'
        doc.styles['Normal'].font.size = Pt(11)

        # Helper function to add paragraph with selective bolding, email formatting, and custom bold segments
        def add_paragraph(text, alignment=WD_ALIGN_PARAGRAPH.LEFT, bold=False, size=11, underline=False, email_addresses=None, bold_segments=None, indent=None):
            email_addresses = email_addresses or []
            bold_segments = bold_segments or []
            pattern_parts = [re.escape(segment) for segment in email_addresses + bold_segments + ['“Party A”', '“Party B”']]
            pattern = r'(' + '|'.join(pattern_parts) + r')' if pattern_parts else r'(“Party A”|“Party B”)'
            paragraphs = text.split('\n\n')
            ps = []
            for para_text in paragraphs:
                p = doc.add_paragraph()
                p.alignment = alignment
                if indent:
                    p.paragraph_format.left_indent = Inches(indent)
                parts = re.split(pattern, para_text)
                for part in parts:
                    run = p.add_run(part)
                    run.font.size = Pt(size)
                    run.bold = bold or part in bold_segments or part in ['“Party A”', '“Party B”']
                    if part in email_addresses:
                        run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color
                        run.underline = WD_UNDERLINE.SINGLE
                    elif underline:
                        run.underline = WD_UNDERLINE.SINGLE
                ps.append(p)
            return ps

        # Helper function to add paragraph with selective bold and size
        def add_paragraph_with_bold(text_parts, bold_parts, alignment=WD_ALIGN_PARAGRAPH.LEFT, default_size=11, bold_size=12, indent=None):
            text = ''.join(text_parts)
            paragraphs = text.split('\n\n')
            ps = []
            for para_text in paragraphs:
                p = doc.add_paragraph()
                p.alignment = alignment
                if indent:
                    p.paragraph_format.left_indent = Inches(indent)
                pattern_parts = [re.escape(bp) for bp in bold_parts] + ['“Party A”', '“Party B”']
                pattern = r'(' + '|'.join(pattern_parts) + r')'
                sub_parts = re.split(pattern, para_text)
                for sub_part in sub_parts:
                    run = p.add_run(sub_part)
                    run.bold = sub_part in bold_parts or sub_part in ['“Party A”', '“Party B”']
                    run.font.size = Pt(bold_size if sub_part in bold_parts else default_size)
                ps.append(p)
            return ps

        # Helper function to add paragraph with selective formatting for Party B email and bold parts
        def add_paragraph_with_email_formatting(text_parts, bold_parts, email_text, alignment=WD_ALIGN_PARAGRAPH.LEFT, default_size=11, bold_size=12):
            text = ''.join(text_parts)
            paragraphs = text.split('\n\n')
            ps = []
            for para_text in paragraphs:
                p = doc.add_paragraph()
                p.alignment = alignment
                # Define pattern for bold_parts (including Party A/B if needed)
                if bold_parts:
                    bold_pattern_parts = [re.escape(bp) for bp in bold_parts] + ['“Party A”', '“Party B”']
                    bold_pattern = r'(' + '|'.join(bold_pattern_parts) + r')'
                else:
                    bold_pattern = r'(“Party A”|“Party B”)'
                # Split on email_text first
                email_parts = para_text.split(email_text)
                for i, email_part in enumerate(email_parts):
                    # Now split each email_part on bold_pattern
                    if bold_pattern:
                        sub_parts = re.split(bold_pattern, email_part)
                    else:
                        sub_parts = [email_part]
                    for sub_part in sub_parts:
                        if sub_part.strip():  # Skip empty strings
                            run = p.add_run(sub_part)
                            is_bold = sub_part in bold_parts or sub_part in ['“Party A”', '“Party B”']
                            run.bold = is_bold
                            run.font.size = Pt(bold_size if is_bold else default_size)
                    # Add the email_text as a separate run if not the last part
                    if i < len(email_parts) - 1:
                        email_run = p.add_run(email_text)
                        email_run.font.size = Pt(default_size)
                        email_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color
                        email_run.underline = WD_UNDERLINE.SINGLE
                ps.append(p)
            return ps

        # Updated Helper function to add heading with 11pt font size
        def add_heading(number, title, level, size=11):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(10)  # Add small space before for separation
            p.paragraph_format.space_after = Pt(0)   # Remove space after for compact layout
            run1 = p.add_run(f"ARTICLE {number}")
            run1.font.name = 'Calibri'
            run1.font.size = Pt(size)
            run1.bold = True
            run1.underline = WD_UNDERLINE.SINGLE
            run1.font.color.rgb = RGBColor(0, 0, 0)
            run2 = p.add_run(": ")
            run2.font.name = 'Calibri'
            run2.font.size = Pt(size)
            run2.bold = True
            run2.font.color.rgb = RGBColor(0, 0, 0)
            run3 = p.add_run(title)
            run3.font.name = 'Calibri'
            run3.font.size = Pt(size)
            run3.bold = True
            run3.font.color.rgb = RGBColor(0, 0, 0)
            return p

        # Define standard articles
        standard_articles = [
            {
                'number': 1,
                'title': 'TERMS OF REFERENCE',
                'content': (
                    '“Party B” shall perform tasks as stated in the attached TOR (annex-1) to “Party A”, '
                    'and deliver each milestone as stipulated in article 4.\n\n'
                    'The work shall be of good quality and well performed with the acceptance by “Party A”.'
                ),
                'table': None
            },
            {
                'number': 2,
                'title': 'TERM OF AGREEMENT',
                'content': (
                    f'The agreement is effective from {contract_data["agreement_start_date_display"]} – '
                    f'{contract_data["agreement_end_date_display"]}. This Agreement is terminated automatically '
                    'after the due date of the Agreement Term unless otherwise, both Parties agree to extend '
                    'the Term with a written agreement.'
                ),
                'table': None
            },
            {
                'number': 3,
                'title': 'PROFESSIONAL FEE',
                'content': [
                    f'The professional fee is the total amount of ',
                    contract_data["total_gross"],
                    f' (',
                    f'{contract_data["total_fee_words"]} ',
                    f') {"excluding" if tax_percentage == 0 else "including"} tax for the whole assignment period.'
                ],
                'financial_lines': [
                    f'{vat_organization_name}' if tax_percentage == 0 and vat_organization_name and deduct_tax_code else '',
                    f'VAT TIN: {deduct_tax_code}' if tax_percentage == 0 and deduct_tax_code else '',
                    f'Total Service Fee: {contract_data["total_gross"]}',
                    f'Withholding Tax {tax_percentage}%: USD{total_gross_amount * (tax_percentage/100):.2f}' if tax_percentage > 0 else '',
                    f'Net amount: {contract_data["total_net"]}',
                ],
                'remaining_content': [
                    f'“Party B” is responsible to issue the Invoice (net amount) and receipt (when receiving the payment) '
                    f'with the total amount as stipulated in each instalment as in the Article 4 after having done the '
                    f'agreed deliverable tasks, for payment request. The payment will be processed after the satisfaction '
                    f'from “Party A” as of the required deliverable tasks as stated in Article 4.\n\n'
                    f'“Party B” is responsible for all related taxes payable to the government department.'
                ],
                'bold_parts': [
                    contract_data["total_gross"],
                    f'{contract_data["total_fee_words"]} ',
                    f'{vat_organization_name}' if tax_percentage == 0 and vat_organization_name and deduct_tax_code else '',
                    f'VAT TIN: {deduct_tax_code}' if tax_percentage == 0 and deduct_tax_code else '',
                    f'Total Service Fee: {contract_data["total_gross"]}',
                    f'Withholding Tax {tax_percentage}%: USD{total_gross_amount * (tax_percentage/100):.2f}' if tax_percentage > 0 else '',
                    f'Net amount: {contract_data["total_net"]}',
                    '“Party A”',
                    '“Party B”'
                ],
                'table': None
            },
            {
                'number': 4,
                'title': 'TERM OF PAYMENT',
                'content': 'The payment will be made based on the following schedules:',
                'table': [
                    {'Installment': 'Installment', 'Total Amount (USD)': ['Total Amount (USD)'], 'Deliverable': 'Deliverable', 'Due date': 'Due date'},
                    *[
                        {
                            'Installment': installment['description'],
                            'Total Amount (USD)': [
                                f'- Gross: ${installment["gross_amount"]:.2f}',
                                f'- Tax {tax_percentage}%: ${installment["tax_amount"]:.2f}' if tax_percentage > 0 else '',
                                f'- Net pay: ${installment["net_amount"]:.2f}'
                            ],
                            'Deliverable': '\n- '.join([d.strip() for d in installment['deliverables'].split(';') if d.strip()]),
                            'Due date': installment['dueDate_display']
                        }
                        for installment in contract_data.get('payment_installments', [])
                    ]
                ]
            },
            {
                'number': 5,
                'title': 'NO OTHER PERSONS',
                'content': (
                    'No person or entity, which is not a party to this agreement, has any rights to enforce, '
                    'take any action, or claim it is owed any benefit under this agreement.'
                ),
                'table': None
            },
            {
                'number': 6,
                'title': 'MONITORING and COORDINATION',
                'content': (
                    f'“Party A” shall monitor and evaluate the progress of the agreement toward its objective, '
                    f'including the activities implemented. '
                    f'{" and ".join([f"{person["name"]}, {person["position"]} (Telephone {person["phone"]} Email: {person["email"]})" for person in contract_data.get("focal_person_info", [])]) or "N/A, N/A (Telephone N/A Email: N/A)"} '
                    f'is the focal contact person of “Party A” and '
                    f'{contract_data.get("party_b_signature_name", "N/A")}, {contract_data.get("party_b_position", "Freelance Consultant")} '
                    f'(HP. {contract_data.get("party_b_phone", "N/A")}, E-mail: {contract_data.get("party_b_email", "N/A")}) '
                    f'the focal contact person of the “Party B”. The focal contact person of “Party A” and “Party B” will work together '
                    f'for overall coordination including reviewing and meeting discussions during the assignment process.'
                ),
                'table': None
            },
            {
                'number': 7,
                'title': 'CONFIDENTIALITY',
                'content': (
                    f'All outputs produced, with the exception of the “{contract_data.get("project_title", "N/A")}”, '
                    f'which is a contribution from, and to be claimed as a public document by the main author and co-author '
                    f'in associated, and/or under this agreement, shall be the property of “Party A”. The “Party B” agrees '
                    f'to not disclose any confidential information, of which he/she may take cognizance in the performance '
                    f'under this contract, except with the prior written approval of “Party A”.'
                ),
                'table': None
            },
            {
                'number': 8,
                'title': 'ANTI-CORRUPTION and CONFLICT OF INTEREST',
                'content': (
                    '“Party B” shall not participate in any practice that is or could be construed as an illegal or corrupt '
                    'practice in Cambodia.\n\nThe “Party A” is committed to fighting all types of corruption and expects this same '
                    'commitment from the consultant. It reserves the rights and believes based on the declaration of “Party B” '
                    'that it is an independent social enterprise firm operating in Cambodia and it does not involve any conflict '
                    'of interest with other parties that may be affected to the “Party A”.'
                ),
                'table': None
            },
            {
                'number': 9,
                'title': 'OBLIGATION TO COMPLY WITH THE NGOF’S POLICIES AND CODE OF CONDUCT',
                'content': (
                    'By signing this agreement, “Party B” is obligated to comply with and respect all existing policies and code '
                    'of conduct of “Party A”, such as Gender Mainstreaming, Child Protection, Disability policy, Environmental '
                    'Mainstreaming, etc. and the “Party B” declared themselves that s/he will perform the assignment in the neutral '
                    'position, professional manner, and not be involved in any political affiliation.'
                ),
                'table': None
            },
            {
                'number': 10,
                'title': 'ANTI-TERRORISM FINANCING AND FINANCIAL CRIME',
                'content': (
                    'NGOF is determined that all its funds and resources should only be used to further its mission and shall not '
                    'be subject to illicit use by any third party nor used or abused for any illicit purpose. In order to achieve '
                    'this objective, NGOF will not knowingly or recklessly provide funds, economic goods, or material support to any '
                    'entity or individual designated as a “terrorist” by the international community or affiliate domestic governments '
                    'and will take all reasonable steps to safeguard and protect its assets from such illicit use and to comply with '
                    'host government laws.\n\n'
                    'NGOF respects its contracts with its donors and puts procedures in place for compliance with these contracts.\n\n'
                    '“Illicit use” refers to terrorist financing, sanctions, money laundering, and export control regulations.'
                ),
                'table': None
            },
            {
                'number': 11,
                'title': 'INSURANCE',
                'content': (
                    '“Party B” is responsible for any health and life insurance of its team members. “Party A” will not be held '
                    'responsible for any medical expenses or compensation incurred during or after this contract.'
                ),
                'table': None
            },
            {
                'number': 12,
                'title': 'ASSIGNMENT',
                'content': (
                    '“Party B” shall have the right to assign individuals within its organization to carry out the tasks herein '
                    'named in the attached Technical Proposal.\n\nThe “Party B” shall not assign, or transfer any of its rights or '
                    'obligations under this agreement without the prior written consent of “Party A”. Any attempt by '
                    '“Party B” to assign or transfer any of its rights and obligations without the prior written consent of “Party A” '
                    'shall render this agreement subject to immediate termination by “Party A”.'
                ),
                'table': None
            },
            {
                'number': 13,
                'title': 'RESOLUTION OF CONFLICTS/DISPUTES',
                'content': (
                    'Conflicts between any of these agreements shall be resolved by the following methods:\n\n'
                    'In the case of a disagreement arising between “Party A” and the “Party B” regarding the implementation of '
                    'any part of, or any other substantive question arising under or relating to this agreement, the parties shall '
                    'use their best efforts to arrive at an agreeable resolution by mutual consultation.\n\n'
                    'Unresolved issues may, upon the option of either party and written notice to the other party, be referred to '
                    'for arbitration. Failure by the “Party B” or “Party A” to dispute a decision arising from such arbitration in '
                    'writing within thirty (30) calendar days of receipt of a final decision shall result in such final decision '
                    'being deemed binding upon either the “Party B” and/or “Party A”. All expenses related to arbitration will be '
                    'shared equally between both parties.'
                ),
                'table': None
            },
            {
                'number': 14,
                'title': 'TERMINATION',
                'content': (
                    'The “Party A” or the “Party B” may, by notice in writing, terminate this agreement under the following conditions:\n\n'
                    '1. “Party A” may terminate this agreement at any time with a one-week notice if “Party B” fails to comply with the '
                    'terms and conditions of this agreement.\n\n'
                    '2. For gross professional misconduct (as defined in the NGOF Human Resource Policy), “Party A” may terminate '
                    'this agreement immediately without prior notice. “Party A” will notify “Party B” in a letter that will indicate '
                    'the reason for termination as well as the effective date of termination.\n\n'
                    '3. “Party B” may terminate this agreement at any time with a one-week notice if “Party A” fails to comply with '
                    'the terms and conditions of this agreement. “Party B” will notify “Party A” in a letter that will indicate the '
                    'reason for termination as well as the effective date of termination. If “Party B” terminates this '
                    'agreement without any appropriate reason or fails to implement the assignment, “Party B” must '
                    'refund the full amount of fees received to “Party A”.\n\n'
                    '4. If for any reason either “Party A” or “Party B” decides to terminate this agreement, “Party B” shall be '
                    'paid pro-rata for the work already completed by “Party A”. This payment will require the submission of a timesheet '
                    'that demonstrates work completed as well as the handing over of any deliverables completed or partially completed. '
                    'In case “Party B” has received payment for services under the agreement which have not yet been performed, the '
                    'appropriate portion of these fees must be refunded by “Party B” to “Party A”.'
                ),
                'table': None
            },
            {
                'number': 15,
                'title': 'MODIFICATION OR AMENDMENT',
                'content': (
                    'No modification or amendment of this agreement shall be valid unless in writing and signed by an authorized '
                    'person of “Party A” and “Party B”.'
                ),
                'table': None
            },
            {
                'number': 16,
                'title': 'CONTROLLING OF LAW',
                'content': (
                    'This agreement shall be governed and construed following the law of the Kingdom of Cambodia. '
                    'This Agreement is prepared in two original copies.'
                ),
                'table': None
            }
        ]

        # Prepare custom articles
        custom_articles = [
            {'article_number': str(k), 'custom_sentence': v}
            for k, v in contract_data.get('custom_article_sentences', {}).items()
        ]

        # Header (adjusted for smaller letterhead space and no underline)
        # Add a paragraph with reduced space before the title for letterhead
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(36)  # Reduced to 0.5 inch (36 points) for a more compact letterhead space
        # Add "The Service Agreement" with no space after
        p = add_paragraph('The Service Agreement', WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, underline=False)[0]
        p.paragraph_format.space_after = Pt(0)  # Remove space after
        # Add "ON" with no space after
        p = add_paragraph('On', WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12)[0]
        p.paragraph_format.space_after = Pt(0)  # Remove space after
        # Add project title (unchanged spacing)
        add_paragraph(contract_data.get('project_title', 'N/A'), WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14)
        # Add contract number (unchanged spacing)
        add_paragraph(f"No.: {contract_data.get('contract_number', 'N/A')}", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14)
        # Add "BETWEEN" (unchanged spacing)
        add_paragraph('BETWEEN', WD_ALIGN_PARAGRAPH.CENTER, size=12)

        # Party A
        party_a_info = contract_data.get('party_a_info', [{'name': 'Mr. SOEUNG Saroeun', 'position': 'Executive Director', 'address': '#9-11, Street 476, Sangkat Tuol Tumpoung I, Phnom Penh, Cambodia'}])
        representatives = [f"{person['name']}, {person['position']}" for person in party_a_info]
        representative_text = ", represented by " + "; ".join(representatives) + "."
        party_a_text_parts = [
            "The NGO Forum on Cambodia",
            representative_text,
            "\nAddress: ",
            party_a_info[0]['address'] if party_a_info else '#9-11, Street 476, Sangkat Tuol Tumpoung I, Phnom Penh, Cambodia',
            ".\nhereinafter called the ",
            "“Party A”"
        ]
        party_a_bold_parts = ["The NGO Forum on Cambodia", "“Party A”"] + [person['name'] for person in party_a_info]
        add_paragraph_with_bold(party_a_text_parts, party_a_bold_parts, WD_ALIGN_PARAGRAPH.CENTER, default_size=12, bold_size=12)

        add_paragraph('AND', WD_ALIGN_PARAGRAPH.CENTER, size=12)

        # Party B
        party_b_position = contract_data.get('party_b_position', 'Freelance Consultant')
        party_b_name = contract_data.get('party_b_signature_name', 'N/A')
        party_b_address = contract_data.get('party_b_address', 'N/A')
        party_b_phone = contract_data.get('party_b_phone', 'N/A')
        party_b_email = contract_data.get('party_b_email', 'N/A')
        party_b_text_parts = [
            party_b_position + " " + party_b_name,
            ",\nAddress: ",
            party_b_address,
            "\nH/P: ",
            party_b_phone,
            ", E-mail: ",
            party_b_email,
            "\nhereinafter called the ",
            "“Party B”"
        ]
        party_b_bold_parts = [party_b_position + " " + party_b_name, "“Party B”"]
        add_paragraph_with_email_formatting(party_b_text_parts, party_b_bold_parts, party_b_email, WD_ALIGN_PARAGRAPH.CENTER, default_size=12, bold_size=12)

        # Whereas Clauses
        add_paragraph(
            f"Whereas NGOF is a legal entity registered with the Ministry of Interior (MOI) "
            f"{contract_data.get('registration_number', '#304 សជណ')} dated {contract_data.get('registration_date', '07 March 2012')}.",
            WD_ALIGN_PARAGRAPH.JUSTIFY, size=11
        )
        add_paragraph(
            f"Whereas NGOF will engage the services of “Party B” which accepts the engagement under the following terms and conditions.",
            WD_ALIGN_PARAGRAPH.JUSTIFY, size=11
        )
        add_paragraph("Both Parties Agreed as follows:", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=11)

        # Articles
        for article in standard_articles:
            add_heading(article['number'], article['title'], level=3, size=11)

            if article['number'] == 3:
                # First part (justified)
                add_paragraph_with_bold(
                    article['content'],
                    article['bold_parts'],
                    WD_ALIGN_PARAGRAPH.JUSTIFY,
                    default_size=11,
                    bold_size=12,
                )
                # Financial lines (left-aligned with indentation, no space after, aligned with tab, labels at 12pt)
                for line in article['financial_lines']:
                    if line:  # Only add non-empty lines
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        p.paragraph_format.left_indent = Inches(0.33)
                        p.paragraph_format.space_after = Pt(0)  # Remove space after paragraph
                        if ':' in line:
                            label, value = line.split(':', 1)
                            p.paragraph_format.tab_stops.add_tab_stop(Inches(2.5))  # Adjust tab stop for better alignment
                            run_label = p.add_run(label + ':')
                            run_label.font.size = Pt(12)  # Set label font size to 12pt
                            run_label.bold = True
                            run_tab = p.add_run('\t')
                            run_value = p.add_run(value.strip())
                            run_value.font.size = Pt(12)  # Match value font size to 12pt
                            run_value.bold = True
                        else:
                            # For VAT line or organization name
                            run = p.add_run(line)
                            run.font.size = Pt(12)  # Match font size to 12pt
                            run.bold = True
                # Remaining content (justified)
                add_paragraph_with_bold(
                    article['remaining_content'],
                    article['bold_parts'],
                    WD_ALIGN_PARAGRAPH.JUSTIFY,
                    default_size=11,
                    bold_size=12
                )
            elif article['number'] == 4:
                add_paragraph(article['content'], WD_ALIGN_PARAGRAPH.JUSTIFY, size=11)
            elif article['number'] == 6:
                email_addresses = [person['email'] for person in contract_data.get("focal_person_info", [])] + [contract_data.get("party_b_email", "N/A")]
                bold_segments = (
                    [f"{person['name']}, {person['position']}" for person in contract_data.get("focal_person_info", [])] +
                    [f"Telephone {person['phone']}" for person in contract_data.get("focal_person_info", [])] +
                    [f"{contract_data.get('party_b_signature_name', 'N/A')}, {contract_data.get('party_b_position', 'Freelance Consultant')}",
                     f"HP. {contract_data.get('party_b_phone', 'N/A')}"]
                )
                add_paragraph(article['content'], WD_ALIGN_PARAGRAPH.JUSTIFY, size=11, email_addresses=email_addresses, bold_segments=bold_segments)
            elif article['number'] == 7:
                bold_segments = [
                    f"“{contract_data.get('project_title', 'N/A')}”"
                ]
                add_paragraph(article['content'], WD_ALIGN_PARAGRAPH.JUSTIFY, size=11, bold_segments=bold_segments)
            else:
                add_paragraph(article['content'], WD_ALIGN_PARAGRAPH.JUSTIFY, size=11)

            if article['table']:
                table = doc.add_table(rows=len(article['table']), cols=len(article['table'][0]))
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.allow_autofit = True

                for row in table.rows:
                    for cell in row.cells:
                        tc = cell._element
                        tcPr = tc.get_or_add_tcPr()
                        for border_name in ['top', 'left', 'bottom', 'right']:
                            border = OxmlElement(f'w:{border_name}')
                            border.set(qn('w:val'), 'single')
                            border.set(qn('w:sz'), '4')
                            border.set(qn('w:color'), '000000')
                            tcPr.append(border)

                for i, row_data in enumerate(article['table']):
                    row_cells = table.rows[i].cells
                    for j, key in enumerate(row_data.keys()):
                        cell = row_cells[j]
                        # Handle 'Total Amount (USD)' as a list of lines
                        if key == 'Total Amount (USD)' and isinstance(row_data[key], list):
                            for line in row_data[key]:
                                if line:  # Only add non-empty lines
                                    p = cell.add_paragraph(line)
                                    p.paragraph_format.space_after = Pt(0)
                                    if i == 0:
                                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    else:
                                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    for run in p.runs:
                                        run.font.size = Pt(12)
                                        run.font.name = 'Calibri'
                                        run.bold = (i == 0) or (i > 0 and key == 'Total Amount (USD)')
                        else:
                            cell.text = row_data[key]
                            for paragraph in cell.paragraphs:
                                paragraph.paragraph_format.space_after = Pt(0)
                                if i == 0:
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                else:
                                    if key == 'Deliverable':
                                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                    else:
                                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                for run in paragraph.runs:
                                    run.font.size = Pt(12)
                                    run.font.name = 'Calibri'
                                    if i == 0:
                                        run.bold = True
                                    elif key == 'Total Amount (USD)':
                                        run.bold = True

            for custom in custom_articles:
                if custom['article_number'] == str(article['number']):
                    add_paragraph(custom['custom_sentence'], WD_ALIGN_PARAGRAPH.JUSTIFY, size=11)
        # ========================
        # SIGNATURE BLOCK
        # ========================

        # Add date centered with space before
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(0)  # remove space after
        run = p.add_run(f"Date: {contract_data.get('agreement_start_date_display', '17th September 2025')}")
        run.bold = True
        run.font.size = Pt(11)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Define tab stop position (move Party B further right)
        tab_position = Inches(5.0)   # adjust between 5.0–5.25 if needed

        # "For" labels row
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(36)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.tab_stops.add_tab_stop(tab_position, WD_TAB_ALIGNMENT.LEFT)

        run = p.add_run('For “Party A”')
        run.bold = True
        run.font.size = Pt(11)
        p.add_run('\t')
        run = p.add_run('For “Party B”')
        run.bold = True
        run.font.size = Pt(11)

        # Signature lines row
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(45)  # enough space for writing
        p.paragraph_format.space_after = Pt(0)    # no extra gap
        p.paragraph_format.tab_stops.add_tab_stop(tab_position, WD_TAB_ALIGNMENT.LEFT)

        p.add_run('__________________')
        p.add_run('\t')
        p.add_run('__________________')

        # Names row
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.tab_stops.add_tab_stop(tab_position, WD_TAB_ALIGNMENT.LEFT)

        run = p.add_run(contract_data.get('party_a_signature_name', 'Mr. SOEUNG Saroeun'))
        run.bold = True
        run.font.size = Pt(11)
        p.add_run('\t')
        run = p.add_run(contract_data.get('party_b_signature_name', 'Mr. Leader Din'))
        run.bold = True
        run.font.size = Pt(11)

        # Positions row
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.tab_stops.add_tab_stop(tab_position, WD_TAB_ALIGNMENT.LEFT)

        signer_position = next(
            (person['position'] for person in party_a_info
            if person['name'] == contract_data.get('party_a_signature_name')),
            'Executive Director'
        )

        run = p.add_run(signer_position)
        run.bold = True
        run.font.size = Pt(11)
        p.add_run('\t')
        run = p.add_run(contract_data.get('party_b_position', 'Freelance Consultant'))
        run.bold = True
        run.font.size = Pt(11)

        # Save the document to a BytesIO stream
        output = BytesIO()
        doc.save(output)
        output.seek(0)

        # Return the file
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f"{sanitize_filename(contract_data.get('party_b_signature_name', 'Contract_' + contract_id))}.docx"
        )

    except Exception as e:
        logger.error(f"Error exporting contract {contract_id} to DOCX: {str(e)}")
        flash("An error occurred while exporting the contract.", 'danger')
        return redirect(url_for('contracts.index'))
    

# Update contract
@contracts_bp.route('/update/<contract_id>', methods=['GET', 'POST'])
@login_required
def update(contract_id):
    contract = Contract.query.get_or_404(contract_id)
    if not current_user.has_role('admin') and contract.user_id != current_user.id:
        flash("You are not authorized to update this contract.", 'danger')
        return redirect(url_for('contracts.index'))
    if contract.deleted_at is not None:
        flash("This contract has been deleted and cannot be updated.", 'danger')
        return redirect(url_for('contracts.index'))

    # Fetch unique Party A data from previous contracts
    previous_contracts = Contract.query.filter(Contract.deleted_at == None).all()
    party_a_data = {}
    for c in previous_contracts:
        for person in c.party_a_info or []:
            if isinstance(person, dict) and person.get('name'):
                name = person['name'].strip()
                normalized_name = name.lower()
                if name and normalized_name not in party_a_data:
                    party_a_data[normalized_name] = {
                        'name': name,
                        'position': person.get('position', '').strip(),
                        'address': person.get('address', '').strip()
                    }

    # Fetch unique Party B data
    party_b_data = {}
    for c in previous_contracts:
        name = c.party_b_signature_name.strip()
        if name and name.lower() not in party_b_data:
            party_b_data[name.lower()] = {
                'original_name': name,
                'position': c.party_b_position or '',
                'phone': c.party_b_phone or '',
                'email': c.party_b_email or '',
                'address': c.party_b_address or ''
            }

    # Fetch unique focal person data
    focal_person_data = {}
    for c in previous_contracts:
        focal_persons = c.focal_person_info or [] 
        for person in focal_persons:
            if isinstance(person, dict) and person.get('name'):
                name = person['name'].strip()
                normalized_name = name.lower()
                if name and normalized_name not in focal_person_data:
                    focal_person_data[normalized_name] = {
                        'name': name,
                        'position': person.get('position', '').strip(),
                        'phone': person.get('phone', '').strip(),
                        'email': person.get('email', '').strip()
                    }

    form_data = {}
    if request.method == 'POST':
        try:
            # Collect simple fields
            party_b_select = request.form.get('party_b_select', '').strip()
            party_b_name = request.form.get('party_b_signature_name', '').strip() if party_b_select == 'new' else party_b_select
            party_a_signer = request.form.get('party_a_signer', '').strip()
            deduct_tax_code = request.form.get('deduct_tax_code', '').strip()
            vat_organization_name = request.form.get('vat_organization_name', '').strip()

            form_data = {
                'project_title': request.form.get('project_title', '').strip(),
                'contract_number': request.form.get('contract_number', '').strip(),
                'output_description': request.form.get('output_description', '').strip(),
                'tax_percentage': float(request.form.get('tax_percentage', '15.0').strip() or 15.0),
                'deduct_tax_code': deduct_tax_code,
                'vat_organization_name': vat_organization_name,
                'organization_name': request.form.get('organization_name', '').strip(),
                'party_b_signature_name': party_b_name,
                'party_b_position': request.form.get('party_b_position', '').strip(),
                'party_b_phone': request.form.get('party_b_phone', '').strip(),
                'party_b_email': request.form.get('party_b_email', '').strip(),
                'party_b_address': request.form.get('party_b_address', '').strip(),
                'agreement_start_date': request.form.get('agreement_start_date', '').strip(),
                'agreement_end_date': request.form.get('agreement_end_date', '').strip(),
                'total_fee_usd': float(request.form.get('total_fee_usd', '0.0').strip() or 0.0),
                'total_fee_words': request.form.get('total_fee_words', '').strip(),
                'workshop_description': request.form.get('workshop_description', '').strip(),
                'title': request.form.get('title', '').strip(),
                'party_b_full_name_with_title': party_b_name,
                'party_b_signature_name_confirm': request.form.get('party_b_signature_name_confirm', '').strip(),
                'party_b_select': party_b_select,
                'party_a_signer': party_a_signer
            }

            # Process Party A info (multiple entries)
            party_a_info = [
                {
                    'name': name.strip(),
                    'position': pos.strip(),
                    'address': addr.strip()
                }
                for name, pos, addr in zip(
                    request.form.getlist('party_a_name[]'),
                    request.form.getlist('party_a_position[]'),
                    request.form.getlist('party_a_address[]')
                )
                if name.strip() and pos.strip() and addr.strip()
            ]
            if not party_a_info:
                flash('At least one Party A representative is required.', 'danger')
                form_data['payment_installments'] = []
                form_data['focal_person_info'] = []
                form_data['articles'] = []
                form_data['party_a_info'] = [{'name': '', 'position': '', 'address': ''}]
                return render_template('contracts/update.html', form_data=form_data, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)

            form_data['party_a_info'] = party_a_info

            # Validate Party A signer
            if not party_a_signer or party_a_signer not in [p['name'] for p in party_a_info]:
                flash('Please select a valid Party A signer from the list.', 'danger')
                form_data['payment_installments'] = []
                form_data['focal_person_info'] = []
                form_data['articles'] = []
                return render_template('contracts/update.html', form_data=form_data, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)

            # Validate Party B name
            if not party_b_name or not re.match(r'^[a-zA-Z\s\.]+$', party_b_name):
                flash('Party B signature name is required and must contain only letters, spaces, and periods.', 'danger')
                form_data['payment_installments'] = []
                form_data['focal_person_info'] = []
                form_data['articles'] = []
                return render_template('contracts/update.html', form_data=form_data, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)

            # Validate vat_organization_name and deduct_tax_code when tax_percentage is 0
            if form_data['tax_percentage'] == 0:
                if not vat_organization_name:
                    flash('VAT Organization Name is required when tax percentage is 0%.', 'danger')
                    form_data['payment_installments'] = []
                    form_data['focal_person_info'] = []
                    form_data['articles'] = []
                    return render_template('contracts/update.html', form_data=form_data, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)
                if not re.match(r'^[a-zA-Z\s\.]+$', vat_organization_name):
                    flash('VAT Organization Name must contain only letters, spaces, and periods.', 'danger')
                    form_data['payment_installments'] = []
                    form_data['focal_person_info'] = []
                    form_data['articles'] = []
                    return render_template('contracts/update.html', form_data=form_data, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)
                if len(vat_organization_name) > 100:
                    flash('VAT Organization Name must not exceed 100 characters.', 'danger')
                    form_data['payment_installments'] = []
                    form_data['focal_person_info'] = []
                    form_data['articles'] = []
                    return render_template('contracts/update.html', form_data=form_data, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)
                if not deduct_tax_code:
                    flash('Deduct TAX Code is required when tax percentage is 0%.', 'danger')
                    form_data['payment_installments'] = []
                    form_data['focal_person_info'] = []
                    form_data['articles'] = []
                    return render_template('contracts/update.html', form_data=form_data, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)
                if not re.match(r'^[A-Z0-9\-]+$', deduct_tax_code):
                    flash('Deduct TAX Code must contain only uppercase letters, numbers, and hyphens.', 'danger')
                    form_data['payment_installments'] = []
                    form_data['focal_person_info'] = []
                    form_data['articles'] = []
                    return render_template('contracts/update.html', form_data=form_data, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)
                if len(deduct_tax_code) > 50:
                    flash('Deduct TAX Code must not exceed 50 characters.', 'danger')
                    form_data['payment_installments'] = []
                    form_data['focal_person_info'] = []
                    form_data['articles'] = []
                    return render_template('contracts/update.html', form_data=form_data, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)

            # Process custom articles
            articles_raw = [
                {'article_number': num.strip(), 'custom_sentence': sent.strip()}
                for num, sent in zip(request.form.getlist('articleNumber[]'), request.form.getlist('customSentence[]'))
                if sent.strip()
            ]
            form_data['articles'] = articles_raw
            form_data['custom_article_sentences'] = {str(article['article_number']): article['custom_sentence'] for article in articles_raw}

            # Process payment installments
            payment_installments_raw = [
                {
                    'description': desc.strip(),
                    'deliverables': deliv.strip(),
                    'dueDate': due.strip()
                }
                for desc, deliv, due in zip(
                    request.form.getlist('paymentInstallmentDesc[]'),
                    request.form.getlist('paymentInstallmentDeliverables[]'),
                    request.form.getlist('paymentInstallmentDueDate[]')
                )
                if desc.strip() and deliv.strip() and due.strip()
            ]
            if not payment_installments_raw:
                flash('At least one payment installment is required.', 'danger')
                form_data['payment_installments'] = []
                form_data['focal_person_info'] = []
                return render_template('contracts/update.html', form_data=form_data, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)

            form_data['payment_installments'] = payment_installments_raw
            deliverables = '; '.join([inst['deliverables'] for inst in payment_installments_raw])
            form_data['deliverables'] = deliverables

            # Process focal persons
            focal_person_raw = [
                {
                    'name': name.strip(),
                    'position': pos.strip(),
                    'phone': phone.strip(),
                    'email': email.strip()
                }
                for name, pos, phone, email in zip(
                    request.form.getlist('focal_person_name[]'),
                    request.form.getlist('focal_person_position[]'),
                    request.form.getlist('focal_person_phone[]'),
                    request.form.getlist('focal_person_email[]')
                )
                if name.strip() and pos.strip() and phone.strip() and email.strip()
            ]
            if not focal_person_raw:
                flash('At least one focal person is required.', 'danger')
                form_data['focal_person_info'] = []
                return render_template('contracts/update.html', form_data=form_data, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)

            form_data['focal_person_info'] = focal_person_raw

            # Calculate payments
            total_fee_usd = form_data['total_fee_usd']
            tax_percentage = form_data['tax_percentage']
            gross_amount_usd = total_fee_usd
            total_gross, total_net = calculate_payments(total_fee_usd, tax_percentage, payment_installments_raw)
            form_data['payment_gross'] = f"${total_gross:.2f} USD"
            form_data['payment_net'] = f"${total_net:.2f} USD"
            form_data['gross_amount_usd'] = gross_amount_usd

            # Validate required fields
            required_fields = [
                ('project_title', 'Project title is required.'),
                ('contract_number', 'Contract number is required.'),
                ('output_description', 'Output description is required.'),
                ('organization_name', 'Organization name is required.'),
                ('party_b_signature_name', 'Party B signature name is required.'),
                ('agreement_start_date', 'Agreement start date is required.'),
                ('agreement_end_date', 'Agreement end date is required.'),
                ('total_fee_usd', 'Total fee USD is required.')
            ]
            for field, message in required_fields:
                if not form_data[field]:
                    flash(message, 'danger')
                    return render_template('contracts/update.html', form_data=form_data, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)

            # Validate Party B confirm match
            if form_data['party_b_signature_name'] != form_data['party_b_signature_name_confirm']:
                flash('Party B signature name confirmation does not match.', 'danger')
                return render_template('contracts/update.html', form_data=form_data, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)

            # Validate contract number format
            if not re.match(r"NGOF/\d{4}-\d{3}", form_data['contract_number']):
                flash('Contract number must follow the format NGOF/YYYY-NNN (e.g., NGOF/2025-005).', 'danger')
                return render_template('contracts/update.html', form_data=form_data, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)

            # Check for duplicate contract number (excluding self)
            existing_contract = Contract.query.filter(
                Contract.contract_number == form_data['contract_number'],
                Contract.id != contract_id,
                Contract.deleted_at == None
            ).first()
            if existing_contract:
                flash('Contract number already exists.', 'danger')
                return render_template('contracts/update.html', form_data=form_data, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)

            # Validate dates
            start_date = form_data['agreement_start_date']
            end_date = form_data['agreement_end_date']
            if start_date and end_date:
                try:
                    if datetime.strptime(end_date, '%Y-%m-%d') < datetime.strptime(start_date, '%Y-%m-%d'):
                        flash('Agreement end date must be after start date.', 'danger')
                        return render_template('contracts/update.html', form_data=form_data, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)
                except ValueError:
                    flash('Invalid date format for agreement start or end date.', 'danger')
                    return render_template('contracts/update.html', form_data=form_data, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)

            # Validate total_fee_usd
            if total_fee_usd < 0:
                flash('Total fee USD cannot be negative.', 'danger')
                return render_template('contracts/update.html', form_data=form_data, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)

            # Validate tax_percentage
            if tax_percentage not in [0, 5, 10, 15, 20]:
                flash('Tax percentage must be one of 0, 5, 10, 15, or 20.', 'danger')
                return render_template('contracts/update.html', form_data=form_data, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)

            # Validate payment installment percentages
            total_percentage = 0.0
            for installment in form_data['payment_installments']:
                match = re.search(r'\((\d+\.?\d*)\%\)', installment['description'])
                if not match:
                    flash(f"Invalid installment description format: {installment['description']}. Must include percentage like (50%).", 'danger')
                    return render_template('contracts/update.html', form_data=form_data, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)
                try:
                    percentage = float(match.group(1))
                    total_percentage += percentage
                except ValueError:
                    flash(f"Invalid percentage in installment description: {installment['description']}.", 'danger')
                    return render_template('contracts/update.html', form_data=form_data, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)
                try:
                    datetime.strptime(installment['dueDate'], '%Y-%m-%d')
                except ValueError:
                    flash(f"Invalid due date for installment: {installment['dueDate']}.", 'danger')
                    return render_template('contracts/update.html', form_data=form_data, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)

            if abs(total_percentage - 100.0) > 0.01:
                flash('Total percentage of payment installments must equal 100%.', 'danger')
                return render_template('contracts/update.html', form_data=form_data, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)

            # Validate focal person info
            for person in form_data['focal_person_info']:
                if not re.match(r'^[a-zA-Z\s\.]+$', person['name']):
                    flash(f"Invalid focal person name: {person['name']}. Only letters, spaces, and periods are allowed.", 'danger')
                    return render_template('contracts/update.html', form_data=form_data, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)
                if not re.match(r'^[a-zA-Z\s]+$', person['position']):
                    flash(f"Invalid focal person position: {person['position']}. Only letters and spaces are allowed.", 'danger')
                    return render_template('contracts/update.html', form_data=form_data, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)
                if not re.match(r'^\+?\d{1,4}([-.\s]?\d{1,4}){2,3}$', person['phone']):
                    flash(f"Invalid focal person phone: {person['phone']}. Use format like 012 845 091, +855 12 845 091, or +85512845091.", 'danger')
                    return render_template('contracts/update.html', form_data=form_data, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)
                if not re.match(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', person['email']):
                    flash(f"Invalid focal person email: {person['email']}.", 'danger')
                    return render_template('contracts/update.html', form_data=form_data, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)

            # Validate Party A info
            for person in form_data['party_a_info']:
                if not re.match(r'^[a-zA-Z\s\.]+$', person['name']):
                    flash(f"Invalid Party A name: {person['name']}. Only letters, spaces, and periods are allowed.", 'danger')
                    return render_template('contracts/update.html', form_data=form_data, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)
                if not re.match(r'^[a-zA-Z\s]+$', person['position']):
                    flash(f"Invalid Party A position: {person['position']}. Only letters and spaces are allowed.", 'danger')
                    return render_template('contracts/update.html', form_data=form_data, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)
                if not person['address']:
                    flash(f"Party A address is required.", 'danger')
                    return render_template('contracts/update.html', form_data=form_data, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)

            # Update contract
            contract.project_title = form_data['project_title']
            contract.contract_number = form_data['contract_number']
            contract.organization_name = form_data['organization_name']
            contract.party_a_info = form_data['party_a_info']
            contract.party_b_full_name_with_title = form_data['party_b_full_name_with_title']
            contract.party_b_address = form_data['party_b_address']
            contract.party_b_phone = form_data['party_b_phone']
            contract.party_b_email = form_data['party_b_email']
            contract.registration_number = '#304 សជណ'
            contract.registration_date = '07 March 2012'
            contract.agreement_start_date = form_data['agreement_start_date']
            contract.agreement_end_date = form_data['agreement_end_date']
            contract.total_fee_usd = form_data['total_fee_usd']
            contract.gross_amount_usd = form_data['gross_amount_usd']
            contract.tax_percentage = form_data['tax_percentage']
            contract.deduct_tax_code = form_data['deduct_tax_code'] if form_data['tax_percentage'] == 0 else None
            contract.vat_organization_name = form_data['vat_organization_name'] if form_data['tax_percentage'] == 0 else None
            contract.payment_gross = form_data['payment_gross']
            contract.payment_net = form_data['payment_net']
            contract.workshop_description = form_data['workshop_description']
            contract.focal_person_info = form_data['focal_person_info']
            contract.party_a_signature_name = form_data['party_a_signer']
            contract.party_b_signature_name = form_data['party_b_signature_name']
            contract.party_b_position = form_data['party_b_position']
            contract.total_fee_words = form_data['total_fee_words'] or number_to_words(form_data['total_fee_usd'])
            contract.title = form_data['title']
            contract.deliverables = form_data['deliverables']
            contract.output_description = form_data['output_description']
            contract.custom_article_sentences = form_data['custom_article_sentences']
            contract.payment_installments = form_data['payment_installments']

            db.session.commit()
            flash('Contract updated successfully!', 'success')
            return redirect(url_for('contracts.index'))
        except Exception as e:
            logger.error(f"Error updating contract: {str(e)}")
            flash("An error occurred while updating the contract.", 'danger')
            return render_template('contracts/update.html', form_data=form_data, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)

    # Initialize form_data for GET request from existing contract
    form_data = contract.to_dict()
    form_data['party_a_info'] = form_data.get('party_a_info') or [{'name': 'Mr. SOEUNG Saroeun', 'position': 'Executive Director', 'address': '#9-11, Street 476, Sangkat Tuol Tumpoung I, Phnom Penh, Cambodia'}]
    form_data['focal_person_info'] = form_data.get('focal_person_info') or [{'name': '', 'position': '', 'phone': '', 'email': ''}]
    form_data['payment_installments'] = form_data.get('payment_installments') or [{'description': '', 'deliverables': '', 'dueDate': ''}]
    form_data['articles'] = [{'article_number': k, 'custom_sentence': v} for k, v in (form_data.get('custom_article_sentences') or {}).items()]
    form_data['custom_article_sentences'] = form_data.get('custom_article_sentences') or {}
    form_data['party_a_signer'] = form_data.get('party_a_signature_name') or 'Mr. SOEUNG Saroeun'
    form_data['deduct_tax_code'] = form_data.get('deduct_tax_code') or ''
    form_data['vat_organization_name'] = form_data.get('vat_organization_name') or ''
    party_b_lower = form_data.get('party_b_signature_name', '').lower().strip()
    form_data['party_b_select'] = next((key for key in party_b_data if key == party_b_lower), 'new')

    return render_template('contracts/update.html', form_data=form_data, party_a_data=party_a_data, party_b_data=party_b_data, focal_person_data=focal_person_data)


# Delete contract
@contracts_bp.route('/delete/<contract_id>', methods=['POST'])
@login_required
def delete(contract_id):
    try:
        contract = Contract.query.get_or_404(contract_id)
        # Allow admins to delete any contract, non-admins only their own
        if not current_user.has_role('admin') and contract.user_id != current_user.id:
            flash("You are not authorized to delete this contract.", 'danger')
            return redirect(url_for('contracts.index'))
        if contract.deleted_at is not None:
            flash("This contract has already been deleted.", 'danger')
            return redirect(url_for('contracts.index'))

        contract.deleted_at = datetime.now()
        db.session.commit()
        flash('Contract deleted successfully!', 'success')
    except Exception as e:
        logger.error(f"Error deleting contract: {str(e)}")
        flash("An error occurred while deleting the contract.", 'danger')
    return redirect(url_for('contracts.index'))

@contracts_bp.route('/export_all_docx', methods=['GET'])
@login_required
def export_all_docx():
    try:
        # Query contracts based on user role
        if current_user.has_role('admin'):
            contracts = Contract.query.filter(Contract.deleted_at == None).all()
        else:
            contracts = Contract.query.filter(Contract.user_id == current_user.id, Contract.deleted_at == None).all()
        
        if not contracts:
            flash("No contracts available to export.", "warning")
            return redirect(url_for('contracts.index'))

        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for contract in contracts:
                try:
                    contract_data = contract.to_dict()
                    if not contract_data:
                        logger.error(f"Contract {contract.id} has no data.")
                        continue

                    # Initialize custom_article_sentences if missing
                    if 'custom_article_sentences' not in contract_data or contract_data['custom_article_sentences'] is None:
                        contract_data['custom_article_sentences'] = {}

                    # Format dates
                    contract_data['agreement_start_date_display'] = format_date(contract_data.get('agreement_start_date', ''))
                    contract_data['agreement_end_date_display'] = format_date(contract_data.get('agreement_end_date', ''))

                    # Get financial data as floats
                    try:
                        total_fee_usd = float(contract_data.get('total_fee_usd', 0.0)) if contract_data.get('total_fee_usd') else 0.0
                        tax_percentage = float(contract_data.get('tax_percentage', 15.0))
                        deduct_tax_code = contract_data.get('deduct_tax_code', '')
                        vat_organization_name = contract_data.get('vat_organization_name', '')  # New field
                    except (ValueError, TypeError) as e:
                        logger.error(f"Error converting financial data for contract {contract.id}: {str(e)}")
                        continue

                    contract_data['total_fee_usd'] = total_fee_usd
                    contract_data['gross_amount_usd'] = total_fee_usd
                    contract_data['total_fee_words'] = contract_data.get('total_fee_words') or number_to_words(total_fee_usd)

                    # Calculate total gross and net
                    total_gross_amount, total_net_amount = calculate_payments(
                        total_fee_usd, tax_percentage, contract_data.get('payment_installments', [])
                    )
                    contract_data['total_gross'] = f"USD{total_gross_amount:.2f}"
                    contract_data['total_net'] = f"USD{total_net_amount:.2f}"

                    # Process payment installments
                    for installment in contract_data.get('payment_installments', []):
                        installment['dueDate_display'] = format_date(installment.get('dueDate', ''))
                        match = re.search(r'\((\d+\.?\d*)\%\)', installment.get('description', ''))
                        percentage = float(match.group(1)) if match else 0.0
                        gross, tax, net = calculate_installment_payments(total_fee_usd, tax_percentage, percentage)
                        installment['gross_amount'] = gross
                        installment['tax_amount'] = tax
                        installment['net_amount'] = net

                    # Create a new DOCX document for each contract
                    doc = Document()

                    # Set document margins (adjusted top margin for the first page letterhead)
                    sections = doc.sections
                    for i, section in enumerate(sections):
                        if i == 0:  # Only modify the first section for letterhead space
                            section.top_margin = Inches(1.5)  # Reduced to 1.5 inches for a more balanced letterhead space
                            section.left_margin = Inches(1)
                            section.right_margin = Inches(1)
                            section.bottom_margin = Inches(1)
                        else:
                            section.top_margin = Inches(1)  # Default margin for subsequent pages
                            section.left_margin = Inches(1)
                            section.right_margin = Inches(1)
                            section.bottom_margin = Inches(1)

                    # Set default font
                    doc.styles['Normal'].font.name = 'Calibri'
                    doc.styles['Normal'].font.size = Pt(11)

                    # Helper function to add paragraph with selective bolding, email formatting, and custom bold segments
                    def add_paragraph(text, alignment=WD_ALIGN_PARAGRAPH.LEFT, bold=False, size=11, underline=False, email_addresses=None, bold_segments=None, indent=None):
                        email_addresses = email_addresses or []
                        bold_segments = bold_segments or []
                        pattern_parts = [re.escape(segment) for segment in email_addresses + bold_segments + ['“Party A”', '“Party B”']]
                        pattern = r'(' + '|'.join(pattern_parts) + r')' if pattern_parts else r'(“Party A”|“Party B”)'
                        paragraphs = text.split('\n\n')
                        ps = []
                        for para_text in paragraphs:
                            p = doc.add_paragraph()
                            p.alignment = alignment
                            if indent:
                                p.paragraph_format.left_indent = Inches(indent)
                            parts = re.split(pattern, para_text)
                            for part in parts:
                                run = p.add_run(part)
                                run.font.size = Pt(size)
                                run.bold = bold or part in bold_segments or part in ['“Party A”', '“Party B”']
                                if part in email_addresses:
                                    run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color
                                    run.underline = WD_UNDERLINE.SINGLE
                                elif underline:
                                    run.underline = WD_UNDERLINE.SINGLE
                            ps.append(p)
                        return ps

                    # Helper function to add paragraph with selective bold and size
                    def add_paragraph_with_bold(text_parts, bold_parts, alignment=WD_ALIGN_PARAGRAPH.LEFT, default_size=11, bold_size=12, indent=None):
                        text = ''.join(text_parts)
                        paragraphs = text.split('\n\n')
                        ps = []
                        for para_text in paragraphs:
                            p = doc.add_paragraph()
                            p.alignment = alignment
                            if indent:
                                p.paragraph_format.left_indent = Inches(indent)
                            pattern_parts = [re.escape(bp) for bp in bold_parts] + ['“Party A”', '“Party B”']
                            pattern = r'(' + '|'.join(pattern_parts) + r')'
                            sub_parts = re.split(pattern, para_text)
                            for sub_part in sub_parts:
                                run = p.add_run(sub_part)
                                run.bold = sub_part in bold_parts or sub_part in ['“Party A”', '“Party B”']
                                run.font.size = Pt(bold_size if sub_part in bold_parts else default_size)
                            ps.append(p)
                        return ps

                    # Helper function to add paragraph with selective formatting for Party B email and bold parts
                    def add_paragraph_with_email_formatting(text_parts, bold_parts, email_text, alignment=WD_ALIGN_PARAGRAPH.LEFT, default_size=11, bold_size=12):
                        text = ''.join(text_parts)
                        paragraphs = text.split('\n\n')
                        ps = []
                        for para_text in paragraphs:
                            p = doc.add_paragraph()
                            p.alignment = alignment
                            # Define pattern for bold_parts (including Party A/B if needed)
                            if bold_parts:
                                bold_pattern_parts = [re.escape(bp) for bp in bold_parts] + ['“Party A”', '“Party B”']
                                bold_pattern = r'(' + '|'.join(bold_pattern_parts) + r')'
                            else:
                                bold_pattern = r'(“Party A”|“Party B”)'
                            # Split on email_text first
                            email_parts = para_text.split(email_text)
                            for i, email_part in enumerate(email_parts):
                                # Now split each email_part on bold_pattern
                                if bold_pattern:
                                    sub_parts = re.split(bold_pattern, email_part)
                                else:
                                    sub_parts = [email_part]
                                for sub_part in sub_parts:
                                    if sub_part.strip():  # Skip empty strings
                                        run = p.add_run(sub_part)
                                        is_bold = sub_part in bold_parts or sub_part in ['“Party A”', '“Party B”']
                                        run.bold = is_bold
                                        run.font.size = Pt(bold_size if is_bold else default_size)
                                # Add the email_text as a separate run if not the last part
                                if i < len(email_parts) - 1:
                                    email_run = p.add_run(email_text)
                                    email_run.font.size = Pt(default_size)
                                    email_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color
                                    email_run.underline = WD_UNDERLINE.SINGLE
                            ps.append(p)
                        return ps

                    # Updated Helper function to add heading with 11pt font size
                    def add_heading(number, title, level, size=11):
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        p.paragraph_format.space_before = Pt(10)  # Add small space before for separation
                        p.paragraph_format.space_after = Pt(0)   # Remove space after for compact layout
                        run1 = p.add_run(f"ARTICLE {number}")
                        run1.font.name = 'Calibri'
                        run1.font.size = Pt(size)
                        run1.bold = True
                        run1.underline = WD_UNDERLINE.SINGLE
                        run1.font.color.rgb = RGBColor(0, 0, 0)
                        run2 = p.add_run(": ")
                        run2.font.name = 'Calibri'
                        run2.font.size = Pt(size)
                        run2.bold = True
                        run2.font.color.rgb = RGBColor(0, 0, 0)
                        run3 = p.add_run(title)
                        run3.font.name = 'Calibri'
                        run3.font.size = Pt(size)
                        run3.bold = True
                        run3.font.color.rgb = RGBColor(0, 0, 0)
                        return p

                    # Define standard articles
                    standard_articles = [
                        {
                            'number': 1,
                            'title': 'TERMS OF REFERENCE',
                            'content': (
                                '“Party B” shall perform tasks as stated in the attached TOR (annex-1) to “Party A”, '
                                'and deliver each milestone as stipulated in article 4.\n\n'
                                'The work shall be of good quality and well performed with the acceptance by “Party A”.'
                            ),
                            'table': None
                        },
                        {
                            'number': 2,
                            'title': 'TERM OF AGREEMENT',
                            'content': (
                                f'The agreement is effective from {contract_data["agreement_start_date_display"]} – '
                                f'{contract_data["agreement_end_date_display"]}. This Agreement is terminated automatically '
                                'after the due date of the Agreement Term unless otherwise, both Parties agree to extend '
                                'the Term with a written agreement.'
                            ),
                            'table': None
                        },
                        {
                            'number': 3,
                            'title': 'PROFESSIONAL FEE',
                            'content': [
                                f'The professional fee is the total amount of ',
                                contract_data["total_gross"],
                                f' (',
                                f'{contract_data["total_fee_words"]} ',
                                f') {"excluding" if tax_percentage == 0 else "including"} tax for the whole assignment period.'
                            ],
                            'financial_lines': [
                                f'{vat_organization_name}' if tax_percentage == 0 and vat_organization_name and deduct_tax_code else '',
                                f'VAT TIN: {deduct_tax_code}' if tax_percentage == 0 and deduct_tax_code else '',
                                f'Total Service Fee: {contract_data["total_gross"]}',
                                f'Withholding Tax {tax_percentage}%: USD{total_gross_amount * (tax_percentage/100):.2f}' if tax_percentage > 0 else '',
                                f'Net amount: {contract_data["total_net"]}',
                            ],
                            'remaining_content': [
                                f'“Party B” is responsible to issue the Invoice (net amount) and receipt (when receiving the payment) '
                                f'with the total amount as stipulated in each instalment as in the Article 4 after having done the '
                                f'agreed deliverable tasks, for payment request. The payment will be processed after the satisfaction '
                                f'from “Party A” as of the required deliverable tasks as stated in Article 4.\n\n'
                                f'“Party B” is responsible for all related taxes payable to the government department.'
                            ],
                            'bold_parts': [
                                contract_data["total_gross"],
                                f'{contract_data["total_fee_words"]} ',
                                f'{vat_organization_name}' if tax_percentage == 0 and vat_organization_name and deduct_tax_code else '',
                                f'VAT TIN: {deduct_tax_code}' if tax_percentage == 0 and deduct_tax_code else '',
                                f'Total Service Fee: {contract_data["total_gross"]}',
                                f'Withholding Tax {tax_percentage}%: USD{total_gross_amount * (tax_percentage/100):.2f}' if tax_percentage > 0 else '',
                                f'Net amount: {contract_data["total_net"]}',
                                '“Party A”',
                                '“Party B”'
                            ],
                            'table': None
                        },
                        {
                            'number': 4,
                            'title': 'TERM OF PAYMENT',
                            'content': 'The payment will be made based on the following schedules:',
                            'table': [
                                {'Installment': 'Installment', 'Total Amount (USD)': ['Total Amount (USD)'], 'Deliverable': 'Deliverable', 'Due date': 'Due date'},
                                *[
                                    {
                                        'Installment': installment['description'],
                                        'Total Amount (USD)': [
                                            f'- Gross: ${installment["gross_amount"]:.2f}',
                                            f'- Tax {tax_percentage}%: ${installment["tax_amount"]:.2f}' if tax_percentage > 0 else '',
                                            f'- Net pay: ${installment["net_amount"]:.2f}'
                                        ],
                                        'Deliverable': '\n- '.join([d.strip() for d in installment['deliverables'].split(';') if d.strip()]),
                                        'Due date': installment['dueDate_display']
                                    }
                                    for installment in contract_data.get('payment_installments', [])
                                ]
                            ]
                        },
                        {
                            'number': 5,
                            'title': 'NO OTHER PERSONS',
                            'content': (
                                'No person or entity, which is not a party to this agreement, has any rights to enforce, '
                                'take any action, or claim it is owed any benefit under this agreement.'
                            ),
                            'table': None
                        },
                        {
                            'number': 6,
                            'title': 'MONITORING and COORDINATION',
                            'content': (
                                f'“Party A” shall monitor and evaluate the progress of the agreement toward its objective, '
                                f'including the activities implemented. '
                                f'{" and ".join([f"{person["name"]}, {person["position"]} (Telephone {person["phone"]} Email: {person["email"]})" for person in contract_data.get("focal_person_info", [])]) or "N/A, N/A (Telephone N/A Email: N/A)"} '
                                f'is the focal contact person of “Party A” and '
                                f'{contract_data.get("party_b_signature_name", "N/A")}, {contract_data.get("party_b_position", "Freelance Consultant")} '
                                f'(HP. {contract_data.get("party_b_phone", "N/A")}, E-mail: {contract_data.get("party_b_email", "N/A")}) '
                                f'the focal contact person of the “Party B”. The focal contact person of “Party A” and “Party B” will work together '
                                f'for overall coordination including reviewing and meeting discussions during the assignment process.'
                            ),
                            'table': None
                        },
                        {
                            'number': 7,
                            'title': 'CONFIDENTIALITY',
                            'content': (
                                f'All outputs produced, with the exception of the “{contract_data.get("project_title", "N/A")}”, '
                                f'which is a contribution from, and to be claimed as a public document by the main author and co-author '
                                f'in associated, and/or under this agreement, shall be the property of “Party A”. The “Party B” agrees '
                                f'to not disclose any confidential information, of which he/she may take cognizance in the performance '
                                f'under this contract, except with the prior written approval of “Party A”.'
                            ),
                            'table': None
                        },
                        {
                            'number': 8,
                            'title': 'ANTI-CORRUPTION and CONFLICT OF INTEREST',
                            'content': (
                                '“Party B” shall not participate in any practice that is or could be construed as an illegal or corrupt '
                                'practice in Cambodia.\n\nThe “Party A” is committed to fighting all types of corruption and expects this same '
                                'commitment from the consultant. It reserves the rights and believes based on the declaration of “Party B” '
                                'that it is an independent social enterprise firm operating in Cambodia and it does not involve any conflict '
                                'of interest with other parties that may be affected to the “Party A”.'
                            ),
                            'table': None
                        },
                        {
                            'number': 9,
                            'title': 'OBLIGATION TO COMPLY WITH THE NGOF’S POLICIES AND CODE OF CONDUCT',
                            'content': (
                                'By signing this agreement, “Party B” is obligated to comply with and respect all existing policies and code '
                                'of conduct of “Party A”, such as Gender Mainstreaming, Child Protection, Disability policy, Environmental '
                                'Mainstreaming, etc. and the “Party B” declared themselves that s/he will perform the assignment in the neutral '
                                'position, professional manner, and not be involved in any political affiliation.'
                            ),
                            'table': None
                        },
                        {
                            'number': 10,
                            'title': 'ANTI-TERRORISM FINANCING AND FINANCIAL CRIME',
                            'content': (
                                'NGOF is determined that all its funds and resources should only be used to further its mission and shall not '
                                'be subject to illicit use by any third party nor used or abused for any illicit purpose. In order to achieve '
                                'this objective, NGOF will not knowingly or recklessly provide funds, economic goods, or material support to any '
                                'entity or individual designated as a “terrorist” by the international community or affiliate domestic governments '
                                'and will take all reasonable steps to safeguard and protect its assets from such illicit use and to comply with '
                                'host government laws.\n\n'
                                'NGOF respects its contracts with its donors and puts procedures in place for compliance with these contracts.\n\n'
                                '“Illicit use” refers to terrorist financing, sanctions, money laundering, and export control regulations.'
                            ),
                            'table': None
                        },
                        {
                            'number': 11,
                            'title': 'INSURANCE',
                            'content': (
                                '“Party B” is responsible for any health and life insurance of its team members. “Party A” will not be held '
                                'responsible for any medical expenses or compensation incurred during or after this contract.'
                            ),
                            'table': None
                        },
                        {
                            'number': 12,
                            'title': 'ASSIGNMENT',
                            'content': (
                                '“Party B” shall have the right to assign individuals within its organization to carry out the tasks herein '
                                'named in the attached Technical Proposal.\n\nThe “Party B” shall not assign, or transfer any of its rights or '
                                'obligations under this agreement without the prior written consent of “Party A”. Any attempt by '
                                '“Party B” to assign or transfer any of its rights and obligations without the prior written consent of “Party A” '
                                'shall render this agreement subject to immediate termination by “Party A”.'
                            ),
                            'table': None
                        },
                        {
                            'number': 13,
                            'title': 'RESOLUTION OF CONFLICTS/DISPUTES',
                            'content': (
                                'Conflicts between any of these agreements shall be resolved by the following methods:\n\n'
                                'In the case of a disagreement arising between “Party A” and the “Party B” regarding the implementation of '
                                'any part of, or any other substantive question arising under or relating to this agreement, the parties shall '
                                'use their best efforts to arrive at an agreeable resolution by mutual consultation.\n\n'
                                'Unresolved issues may, upon the option of either party and written notice to the other party, be referred to '
                                'for arbitration. Failure by the “Party B” or “Party A” to dispute a decision arising from such arbitration in '
                                'writing within thirty (30) calendar days of receipt of a final decision shall result in such final decision '
                                'being deemed binding upon either the “Party B” and/or “Party A”. All expenses related to arbitration will be '
                                'shared equally between both parties.'
                            ),
                            'table': None
                        },
                        {
                            'number': 14,
                            'title': 'TERMINATION',
                            'content': (
                                'The “Party A” or the “Party B” may, by notice in writing, terminate this agreement under the following conditions:\n\n'
                                '1. “Party A” may terminate this agreement at any time with a one-week notice if “Party B” fails to comply with the '
                                'terms and conditions of this agreement.\n\n'
                                '2. For gross professional misconduct (as defined in the NGOF Human Resource Policy), “Party A” may terminate '
                                'this agreement immediately without prior notice. “Party A” will notify “Party B” in a letter that will indicate '
                                'the reason for termination as well as the effective date of termination.\n\n'
                                '3. “Party B” may terminate this agreement at any time with a one-week notice if “Party A” fails to comply with '
                                'the terms and conditions of this agreement. “Party B” will notify “Party A” in a letter that will indicate the '
                                'reason for termination as well as the effective date of termination. If “Party B” terminates this '
                                'agreement without any appropriate reason or fails to implement the assignment, “Party B” must '
                                'refund the full amount of fees received to “Party A”.\n\n'
                                '4. If for any reason either “Party A” or “Party B” decides to terminate this agreement, “Party B” shall be '
                                'paid pro-rata for the work already completed by “Party A”. This payment will require the submission of a timesheet '
                                'that demonstrates work completed as well as the handing over of any deliverables completed or partially completed. '
                                'In case “Party B” has received payment for services under the agreement which have not yet been performed, the '
                                'appropriate portion of these fees must be refunded by “Party B” to “Party A”.'
                            ),
                            'table': None
                        },
                        {
                            'number': 15,
                            'title': 'MODIFICATION OR AMENDMENT',
                            'content': (
                                'No modification or amendment of this agreement shall be valid unless in writing and signed by an authorized '
                                'person of “Party A” and “Party B”.'
                            ),
                            'table': None
                        },
                        {
                            'number': 16,
                            'title': 'CONTROLLING OF LAW',
                            'content': (
                                'This agreement shall be governed and construed following the law of the Kingdom of Cambodia. '
                                'This Agreement is prepared in two original copies.'
                            ),
                            'table': None
                        }
                    ]

                    # Prepare custom articles
                    custom_articles = [
                        {'article_number': str(k), 'custom_sentence': v}
                        for k, v in contract_data.get('custom_article_sentences', {}).items()
                    ]

                    # Header (adjusted for smaller letterhead space and no underline)
                    # Add a paragraph with reduced space before the title for letterhead
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(36)  # Reduced to 0.5 inch (36 points) for a more compact letterhead space
                    # Add "The Service Agreement" with no space after
                    p = add_paragraph('The Service Agreement', WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, underline=False)[0]
                    p.paragraph_format.space_after = Pt(0)  # Remove space after
                    # Add "ON" with no space after
                    p = add_paragraph('On', WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12)[0]
                    p.paragraph_format.space_after = Pt(0)  # Remove space after
                    # Add project title (unchanged spacing)
                    add_paragraph(contract_data.get('project_title', 'N/A'), WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14)
                    # Add contract number (unchanged spacing)
                    add_paragraph(f"No.: {contract_data.get('contract_number', 'N/A')}", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14)
                    # Add "BETWEEN" (unchanged spacing)
                    add_paragraph('BETWEEN', WD_ALIGN_PARAGRAPH.CENTER, size=12)

                    # Party A
                    party_a_info = contract_data.get('party_a_info', [{'name': 'Mr. SOEUNG Saroeun', 'position': 'Executive Director', 'address': '#9-11, Street 476, Sangkat Tuol Tumpoung I, Phnom Penh, Cambodia'}])
                    representatives = [f"{person['name']}, {person['position']}" for person in party_a_info]
                    representative_text = ", represented by " + "; ".join(representatives) + "."
                    party_a_text_parts = [
                        "The NGO Forum on Cambodia",
                        representative_text,
                        "\nAddress: ",
                        party_a_info[0]['address'] if party_a_info else '#9-11, Street 476, Sangkat Tuol Tumpoung I, Phnom Penh, Cambodia',
                        ".\nhereinafter called the ",
                        "“Party A”"
                    ]
                    party_a_bold_parts = ["The NGO Forum on Cambodia", "“Party A”"] + [person['name'] for person in party_a_info]
                    add_paragraph_with_bold(party_a_text_parts, party_a_bold_parts, WD_ALIGN_PARAGRAPH.CENTER, default_size=12, bold_size=12)

                    add_paragraph('AND', WD_ALIGN_PARAGRAPH.CENTER, size=12)

                    # Party B
                    party_b_position = contract_data.get('party_b_position', 'Freelance Consultant')
                    party_b_name = contract_data.get('party_b_signature_name', 'N/A')
                    party_b_address = contract_data.get('party_b_address', 'N/A')
                    party_b_phone = contract_data.get('party_b_phone', 'N/A')
                    party_b_email = contract_data.get('party_b_email', 'N/A')
                    party_b_text_parts = [
                        party_b_position + " " + party_b_name,
                        ",\nAddress: ",
                        party_b_address,
                        "\nH/P: ",
                        party_b_phone,
                        ", E-mail: ",
                        party_b_email,
                        "\nhereinafter called the ",
                        "“Party B”"
                    ]
                    party_b_bold_parts = [party_b_position + " " + party_b_name, "“Party B”"]
                    add_paragraph_with_email_formatting(party_b_text_parts, party_b_bold_parts, party_b_email, WD_ALIGN_PARAGRAPH.CENTER, default_size=12, bold_size=12)

                    # Whereas Clauses
                    add_paragraph(
                        f"Whereas NGOF is a legal entity registered with the Ministry of Interior (MOI) "
                        f"{contract_data.get('registration_number', '#304 សជណ')} dated {contract_data.get('registration_date', '07 March 2012')}.",
                        WD_ALIGN_PARAGRAPH.JUSTIFY, size=11
                    )
                    add_paragraph(
                        f"Whereas NGOF will engage the services of “Party B” which accepts the engagement under the following terms and conditions.",
                        WD_ALIGN_PARAGRAPH.JUSTIFY, size=11
                    )
                    add_paragraph("Both Parties Agreed as follows:", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=11)

                    # Articles
                    for article in standard_articles:
                        add_heading(article['number'], article['title'], level=3, size=11)

                        if article['number'] == 3:
                            # First part (justified)
                            add_paragraph_with_bold(
                                article['content'],
                                article['bold_parts'],
                                WD_ALIGN_PARAGRAPH.JUSTIFY,
                                default_size=11,
                                bold_size=12,
                            )
                            # Financial lines (left-aligned with indentation, no space after, aligned with tab, labels at 12pt)
                            for line in article['financial_lines']:
                                if line:  # Only add non-empty lines
                                    p = doc.add_paragraph()
                                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                    p.paragraph_format.left_indent = Inches(0.33)
                                    p.paragraph_format.space_after = Pt(0)  # Remove space after paragraph
                                    if ':' in line:
                                        label, value = line.split(':', 1)
                                        p.paragraph_format.tab_stops.add_tab_stop(Inches(2.5))  # Adjust tab stop for better alignment
                                        run_label = p.add_run(label + ':')
                                        run_label.font.size = Pt(12)  # Set label font size to 12pt
                                        run_label.bold = True
                                        run_tab = p.add_run('\t')
                                        run_value = p.add_run(value.strip())
                                        run_value.font.size = Pt(12)  # Match value font size to 12pt
                                        run_value.bold = True
                                    else:
                                        # For VAT line or organization name
                                        run = p.add_run(line)
                                        run.font.size = Pt(12)  # Match font size to 12pt
                                        run.bold = True
                            # Remaining content (justified)
                            add_paragraph_with_bold(
                                article['remaining_content'],
                                article['bold_parts'],
                                WD_ALIGN_PARAGRAPH.JUSTIFY,
                                default_size=11,
                                bold_size=12
                            )
                        elif article['number'] == 4:
                            add_paragraph(article['content'], WD_ALIGN_PARAGRAPH.JUSTIFY, size=11)
                        elif article['number'] == 6:
                            email_addresses = [person['email'] for person in contract_data.get("focal_person_info", [])] + [contract_data.get("party_b_email", "N/A")]
                            bold_segments = (
                                [f"{person['name']}, {person['position']}" for person in contract_data.get("focal_person_info", [])] +
                                [f"Telephone {person['phone']}" for person in contract_data.get("focal_person_info", [])] +
                                [f"{contract_data.get('party_b_signature_name', 'N/A')}, {contract_data.get('party_b_position', 'Freelance Consultant')}",
                                 f"HP. {contract_data.get('party_b_phone', 'N/A')}"]
                            )
                            add_paragraph(article['content'], WD_ALIGN_PARAGRAPH.JUSTIFY, size=11, email_addresses=email_addresses, bold_segments=bold_segments)
                        elif article['number'] == 7:
                            bold_segments = [
                                f"“{contract_data.get('project_title', 'N/A')}”"
                            ]
                            add_paragraph(article['content'], WD_ALIGN_PARAGRAPH.JUSTIFY, size=11, bold_segments=bold_segments)
                        else:
                            add_paragraph(article['content'], WD_ALIGN_PARAGRAPH.JUSTIFY, size=11)

                        if article['table']:
                            table = doc.add_table(rows=len(article['table']), cols=len(article['table'][0]))
                            table.alignment = WD_TABLE_ALIGNMENT.CENTER
                            table.allow_autofit = True

                            for row in table.rows:
                                for cell in row.cells:
                                    tc = cell._element
                                    tcPr = tc.get_or_add_tcPr()
                                    for border_name in ['top', 'left', 'bottom', 'right']:
                                        border = OxmlElement(f'w:{border_name}')
                                        border.set(qn('w:val'), 'single')
                                        border.set(qn('w:sz'), '4')
                                        border.set(qn('w:color'), '000000')
                                        tcPr.append(border)

                            for i, row_data in enumerate(article['table']):
                                row_cells = table.rows[i].cells
                                for j, key in enumerate(row_data.keys()):
                                    cell = row_cells[j]
                                    # Handle 'Total Amount (USD)' as a list of lines
                                    if key == 'Total Amount (USD)' and isinstance(row_data[key], list):
                                        for line in row_data[key]:
                                            if line:  # Only add non-empty lines
                                                p = cell.add_paragraph(line)
                                                p.paragraph_format.space_after = Pt(0)
                                                if i == 0:
                                                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                                else:
                                                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                                for run in p.runs:
                                                    run.font.size = Pt(12)
                                                    run.font.name = 'Calibri'
                                                    run.bold = (i == 0) or (i > 0 and key == 'Total Amount (USD)')
                                    else:
                                        cell.text = row_data[key]
                                        for paragraph in cell.paragraphs:
                                            paragraph.paragraph_format.space_after = Pt(0)
                                            if i == 0:
                                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                            else:
                                                if key == 'Deliverable':
                                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                                else:
                                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                            for run in paragraph.runs:
                                                run.font.size = Pt(12)
                                                run.font.name = 'Calibri'
                                                if i == 0:
                                                    run.bold = True
                                                elif key == 'Total Amount (USD)':
                                                    run.bold = True

                        for custom in custom_articles:
                            if custom['article_number'] == str(article['number']):
                                add_paragraph(custom['custom_sentence'], WD_ALIGN_PARAGRAPH.JUSTIFY, size=11)
                    # ========================
                    # SIGNATURE BLOCK
                    # ========================

                    # Add date centered with space before
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(24)
                    p.paragraph_format.space_after = Pt(0)  # remove space after
                    run = p.add_run(f"Date: {contract_data.get('agreement_start_date_display', 'N/A')}")
                    run.bold = True
                    run.font.size = Pt(11)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Define tab stop position (move Party B further right)
                    tab_position = Inches(5.0)   # adjust between 5.0–5.25 if needed

                    # "For" labels row
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(36)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.tab_stops.add_tab_stop(tab_position, WD_TAB_ALIGNMENT.LEFT)

                    run = p.add_run('For “Party A”')
                    run.bold = True
                    run.font.size = Pt(11)
                    p.add_run('\t')
                    run = p.add_run('For “Party B”')
                    run.bold = True
                    run.font.size = Pt(11)

                    # Signature lines row
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(45)  # enough space for writing
                    p.paragraph_format.space_after = Pt(0)    # no extra gap
                    p.paragraph_format.tab_stops.add_tab_stop(tab_position, WD_TAB_ALIGNMENT.LEFT)

                    p.add_run('__________________')
                    p.add_run('\t')
                    p.add_run('__________________')

                    # Names row
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.tab_stops.add_tab_stop(tab_position, WD_TAB_ALIGNMENT.LEFT)

                    run = p.add_run(contract_data.get('party_a_signature_name', 'Mr. SOEUNG Saroeun'))
                    run.bold = True
                    run.font.size = Pt(11)
                    p.add_run('\t')
                    run = p.add_run(contract_data.get('party_b_signature_name', 'N/A'))
                    run.bold = True
                    run.font.size = Pt(11)

                    # Positions row
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.tab_stops.add_tab_stop(tab_position, WD_TAB_ALIGNMENT.LEFT)

                    signer_position = next(
                        (person['position'] for person in party_a_info
                        if person['name'] == contract_data.get('party_a_signature_name')),
                        'Executive Director'
                    )

                    run = p.add_run(signer_position)
                    run.bold = True
                    run.font.size = Pt(11)
                    p.add_run('\t')
                    run = p.add_run(contract_data.get('party_b_position', 'Freelance Consultant'))
                    run.bold = True
                    run.font.size = Pt(11)

                    # Save document to buffer
                    doc_buffer = BytesIO()
                    doc.save(doc_buffer)
                    doc_buffer.seek(0)

                    # Generate filename
                    party_b_name = contract_data.get('party_b_signature_name', f'Contract_{contract.id}')
                    filename = f"{sanitize_filename(party_b_name)}.docx"
                    zip_file.writestr(filename, doc_buffer.getvalue())
                    doc_buffer.close()

                except Exception as e:
                    logger.error(f"Error processing contract {contract.id}: {str(e)}")
                    continue

        zip_buffer.seek(0)
        return send_file(
            zip_buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name="All_Contracts.zip"
        )

    except Exception as e:
        logger.error(f"Error exporting all contracts to ZIP: {str(e)}")
        flash("An error occurred while exporting all contracts.", 'danger')
        return redirect(url_for('contracts.index'))
    
# Export contract to excel (original, user-specific)
@contracts_bp.route('/export_excel')
@login_required
def export_excel():
    try:
        search_query = request.args.get('search', '', type=str)
        sort_order = request.args.get('sort', 'created_at_desc', type=str)

        # Filter by user_id and exclude soft-deleted contracts
        query = Contract.query.filter(Contract.user_id == current_user.id, Contract.deleted_at == None)

        # Apply search filter
        if search_query:
            query = query.filter(
                (Contract.project_title.ilike(f'%{search_query}%')) |
                (Contract.contract_number.ilike(f'%{search_query}%')) |
                (Contract.party_b_signature_name.ilike(f'%{search_query}%'))
            )

        # Sorting
        if sort_order == 'contract_number_asc':
            query = query.order_by(Contract.contract_number.asc())
        elif sort_order == 'contract_number_desc':
            query = query.order_by(Contract.contract_number.desc())
        elif sort_order == 'start_date_asc':
            query = query.order_by(Contract.agreement_start_date.asc())
        elif sort_order == 'start_date_desc':
            query = query.order_by(Contract.agreement_start_date.desc())
        elif sort_order == 'total_fee_asc':
            query = query.order_by(Contract.total_fee_usd.asc())
        elif sort_order == 'total_fee_desc':
            query = query.order_by(Contract.total_fee_usd.desc())
        else:  # Default to created_at_desc
            query = query.order_by(Contract.created_at.desc())

        contracts = [contract.to_dict() for contract in query.all()]
        data = []

        for contract in contracts:
            total_fee_usd = float(contract['total_fee_usd']) if contract['total_fee_usd'] else 0.0
            tax_percentage = float(contract.get('tax_percentage', 15.0))
            if contract.get('project_title') == 'REJECTED':
                continue

            # Use the actual contract_number from the database
            formatted_contract_no = contract.get('contract_number', '')

            payment_installments = contract.get('payment_installments', [])
            for idx, installment in enumerate(payment_installments, 1):
                match = re.search(r'\((\d+\.?\d*)\%\)', installment['description'])
                percentage = float(match.group(1)) if match else 0.0
                due_date = format_date(installment.get('dueDate', ''))
                gross, tax, net = calculate_installment_payments(total_fee_usd, tax_percentage, percentage) if match else (0.0, 0.0, 0.0)
                payment_details = (
                    f"Gross: {gross:.2f} USD\n"
                    f"Tax({tax_percentage:.1f}%): {tax:.2f} USD\n"
                    f"Net: {net:.2f} USD"
                )
                data.append({
                    'Contract No.': formatted_contract_no,
                    'Consultant': contract['party_b_signature_name'] or '',
                    'Agreement Name': contract['project_title'] or '',
                    'Term of Payment': f"Installment #{idx} ({percentage:.1f}%)" if percentage else installment['description'],
                    'Date': due_date,
                    '': payment_details,
                    'Attached': ''
                })
            # Empty separator row
            data.append({
                'Contract No.': '',
                'Consultant': '',
                'Agreement Name': '',
                'Term of Payment': '',
                'Date': '',
                '': '',
                'Attached': ''
            })

        df = pd.DataFrame(data)
        output = BytesIO()
        wb = Workbook()
        ws = wb.active
        ws.title = 'List'

        # Row 1: default (no fill)
        ws.row_dimensions[1].height = 5

        # Header row (row 2)
        headers = ['Contract No.', 'Consultant', 'Agreement Name', 'Term of Payment', 'Attached']
        for col_num, header in enumerate(headers, 1):
            target_col = col_num if col_num <= 3 else 4 if col_num == 4 else 7
            cell = ws.cell(row=2, column=target_col, value=header)
            cell.fill = PatternFill(start_color="88B84D", end_color="88B84D", fill_type="solid")
            cell.font = Font(name="Times New Roman", bold=True, size=16)
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell.border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin', color='000000')
            )
        ws.merge_cells(start_row=2, start_column=4, end_row=2, end_column=6)
        ws.cell(row=2, column=4).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        ws.cell(row=2, column=4).border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin', color='000000')
        )
        ws.cell(row=2, column=4).fill = PatternFill(start_color="88B84D", end_color="88B84D", fill_type="solid")

        # Empty teal row UNDER headers (row 3)
        for col in range(1, 8):
            cell = ws.cell(row=3, column=col, value="")
            cell.fill = PatternFill(start_color="28677A", end_color="28677A", fill_type="solid")
            cell.border = Border()
        ws.row_dimensions[3].height = 5

        # Write data rows (start at row 4)
        for r_idx, row in enumerate(dataframe_to_rows(df, index=False, header=False), 4):
            is_separator_row = all(v == "" for v in row)
            for c_idx, value in enumerate(row, 1):
                cell = ws.cell(row=r_idx, column=c_idx, value=value)

                if not is_separator_row:
                    if c_idx in [4, 5, 6]:
                        cell.font = Font(name="Times New Roman", size=14, bold=True, color='FF0000' if c_idx == 6 else '000000')
                    else:
                        cell.font = Font(name="Times New Roman", size=14)

                    cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                    cell.border = Border(
                        left=Side(style='thin'),
                        right=Side(style='thin'),
                        top=Side(style='thin'),
                        bottom=Side(style='thin')
                    )

                    if c_idx in [6, 7]:
                        ws.row_dimensions[r_idx].height = 60
                else:
                    for col in range(1, 8):
                        ws.cell(row=r_idx, column=col, value="")
                        ws.cell(row=r_idx, column=col).fill = PatternFill(start_color="28677A", end_color="28677A", fill_type="solid")
                        ws.cell(row=r_idx, column=col).border = Border()
                    ws.row_dimensions[r_idx].height = 5

        # Merge contract info cells
        current_contract = None
        start_row = 4
        for idx, row in enumerate(data, 4):
            if row['Contract No.'] == '' and current_contract is not None:
                if idx - 1 > start_row:
                    ws.merge_cells(start_row=start_row, start_column=1, end_row=idx-1, end_column=1)
                    ws.merge_cells(start_row=start_row, start_column=2, end_row=idx-1, end_column=2)
                    ws.merge_cells(start_row=start_row, start_column=3, end_row=idx-1, end_column=3)
                    for col in [1, 2, 3]:
                        ws.cell(row=start_row, column=col).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                current_contract = None
                start_row = idx + 1
            elif row['Contract No.'] and current_contract != row['Contract No.']:
                current_contract = row['Contract No.']
                start_row = idx
        if current_contract is not None and len(data) + 3 > start_row:
            ws.merge_cells(start_row=start_row, start_column=1, end_row=len(data)+3, end_column=1)
            ws.merge_cells(start_row=start_row, start_column=2, end_row=len(data)+3, end_column=2)
            ws.merge_cells(start_row=start_row, start_column=3, end_row=len(data)+3, end_column=3)
            for col in [1, 2, 3]:
                ws.cell(row=start_row, column=col).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

        # Column widths
        column_widths = [22, 22, 60, 22, 22, 30, 25]
        for i, width in enumerate(column_widths, 1):
            ws.column_dimensions[chr(64 + i)].width = width

        wb.save(output)
        output.seek(0)

        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name='Consultancy_Agreement_List.xlsx'
        )

    except Exception as e:
        logger.error(f"Error exporting to Excel: {str(e)}")
        flash("An error occurred while exporting to Excel.", 'danger')
        return redirect(url_for('contracts.index'))

# Export all contracts to excel (admin only)
@contracts_bp.route('/export_excel_all')
@login_required
def export_excel_all():
    if not current_user.has_role('admin'):
        flash("You do not have permission to export all contracts.", 'danger')
        return redirect(url_for('contracts.index'))

    try:
        search_query = request.args.get('search', '', type=str)
        sort_order = request.args.get('sort', 'created_at_desc', type=str)

        # Base query for all contracts (exclude soft-deleted)
        query = Contract.query.filter(Contract.deleted_at == None)

        # Apply search filter
        if search_query:
            query = query.filter(
                (Contract.project_title.ilike(f'%{search_query}%')) |
                (Contract.contract_number.ilike(f'%{search_query}%')) |
                (Contract.party_b_signature_name.ilike(f'%{search_query}%'))
            )

        # Sorting
        if sort_order == 'contract_number_asc':
            query = query.order_by(Contract.contract_number.asc())
        elif sort_order == 'contract_number_desc':
            query = query.order_by(Contract.contract_number.desc())
        elif sort_order == 'start_date_asc':
            query = query.order_by(Contract.agreement_start_date.asc())
        elif sort_order == 'start_date_desc':
            query = query.order_by(Contract.agreement_start_date.desc())
        elif sort_order == 'total_fee_asc':
            query = query.order_by(Contract.total_fee_usd.asc())
        elif sort_order == 'total_fee_desc':
            query = query.order_by(Contract.total_fee_usd.desc())
        else:  # Default to created_at_desc
            query = query.order_by(Contract.created_at.desc())

        contracts = [contract.to_dict() for contract in query.all()]
        if not contracts:
            flash("No contracts available to export.", 'warning')
            return redirect(url_for('contracts.index'))

        data = []
        for contract in contracts:
            total_fee_usd = float(contract.get('total_fee_usd', 0.0)) if contract.get('total_fee_usd') is not None else 0.0
            tax_percentage = float(contract.get('tax_percentage', 15.0))
            if contract.get('project_title') == 'REJECTED':
                continue

            # Use the actual contract_number from the database
            formatted_contract_no = contract.get('contract_number', '')

            payment_installments = contract.get('payment_installments', []) or []
            for idx, installment in enumerate(payment_installments, 1):
                match = re.search(r'\((\d+\.?\d*)\%\)', installment.get('description', ''))
                percentage = float(match.group(1)) if match else 0.0
                due_date = format_date(installment.get('dueDate', ''))
                gross, tax, net = calculate_installment_payments(total_fee_usd, tax_percentage, percentage) if match else (0.0, 0.0, 0.0)
                payment_details = (
                    f"Gross: {gross:.2f} USD\n"
                    f"Tax({tax_percentage:.1f}%): {tax:.2f} USD\n"
                    f"Net: {net:.2f} USD"
                )
                data.append({
                    'Contract No.': formatted_contract_no,
                    'Consultant': contract.get('party_b_signature_name', '') or '',
                    'Agreement Name': contract.get('project_title', '') or '',
                    'Term of Payment': f"Installment #{idx} ({percentage:.1f}%)" if percentage else installment.get('description', ''),
                    'Date': due_date,
                    '': payment_details,
                    'Attached': ''
                })
            # Empty separator row
            data.append({
                'Contract No.': '',
                'Consultant': '',
                'Agreement Name': '',
                'Term of Payment': '',
                'Date': '',
                '': '',
                'Attached': ''
            })

        df = pd.DataFrame(data)
        output = BytesIO()
        wb = Workbook()
        ws = wb.active
        ws.title = 'List'

        # Row 1: default (no fill)
        ws.row_dimensions[1].height = 5

        # Header row (row 2)
        headers = ['Contract No.', 'Consultant', 'Agreement Name', 'Term of Payment', 'Attached']
        for col_num, header in enumerate(headers, 1):
            target_col = col_num if col_num <= 3 else 4 if col_num == 4 else 7
            cell = ws.cell(row=2, column=target_col, value=header)
            cell.fill = PatternFill(start_color="88B84D", end_color="88B84D", fill_type="solid")
            cell.font = Font(name="Times New Roman", bold=True, size=16)
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell.border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin', color='000000')
            )
        ws.merge_cells(start_row=2, start_column=4, end_row=2, end_column=6)
        ws.cell(row=2, column=4).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        ws.cell(row=2, column=4).border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin', color='000000')
        )
        ws.cell(row=2, column=4).fill = PatternFill(start_color="88B84D", end_color="88B84D", fill_type="solid")

        # Empty teal row UNDER headers (row 3)
        for col in range(1, 8):
            cell = ws.cell(row=3, column=col, value="")
            cell.fill = PatternFill(start_color="28677A", end_color="28677A", fill_type="solid")
            cell.border = Border()
        ws.row_dimensions[3].height = 5

        # Write data rows (start at row 4)
        for r_idx, row in enumerate(dataframe_to_rows(df, index=False, header=False), 4):
            is_separator_row = all(v == "" for v in row)
            for c_idx, value in enumerate(row, 1):
                cell = ws.cell(row=r_idx, column=c_idx, value=value)

                if not is_separator_row:
                    if c_idx in [4, 5, 6]:
                        cell.font = Font(name="Times New Roman", size=14, bold=True, color='FF0000' if c_idx == 6 else '000000')
                    else:
                        cell.font = Font(name="Times New Roman", size=14)

                    cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                    cell.border = Border(
                        left=Side(style='thin'),
                        right=Side(style='thin'),
                        top=Side(style='thin'),
                        bottom=Side(style='thin')
                    )

                    if c_idx in [6, 7]:
                        ws.row_dimensions[r_idx].height = 60
                else:
                    for col in range(1, 8):
                        ws.cell(row=r_idx, column=col, value="")
                        ws.cell(row=r_idx, column=col).fill = PatternFill(start_color="28677A", end_color="28677A", fill_type="solid")
                        ws.cell(row=r_idx, column=col).border = Border()
                    ws.row_dimensions[r_idx].height = 5

        # Merge contract info cells
        current_contract = None
        start_row = 4
        for idx, row in enumerate(data, 4):
            if row['Contract No.'] == '' and current_contract is not None:
                if idx - 1 > start_row:
                    ws.merge_cells(start_row=start_row, start_column=1, end_row=idx-1, end_column=1)
                    ws.merge_cells(start_row=start_row, start_column=2, end_row=idx-1, end_column=2)
                    ws.merge_cells(start_row=start_row, start_column=3, end_row=idx-1, end_column=3)
                    for col in [1, 2, 3]:
                        ws.cell(row=start_row, column=col).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                current_contract = None
                start_row = idx + 1
            elif row['Contract No.'] and current_contract != row['Contract No.']:
                current_contract = row['Contract No.']
                start_row = idx
        if current_contract is not None and len(data) + 3 > start_row:
            ws.merge_cells(start_row=start_row, start_column=1, end_row=len(data)+3, end_column=1)
            ws.merge_cells(start_row=start_row, start_column=2, end_row=len(data)+3, end_column=2)
            ws.merge_cells(start_row=start_row, start_column=3, end_row=len(data)+3, end_column=3)
            for col in [1, 2, 3]:
                ws.cell(row=start_row, column=col).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

        # Column widths
        column_widths = [22, 22, 60, 22, 22, 30, 25]
        for i, width in enumerate(column_widths, 1):
            ws.column_dimensions[chr(64 + i)].width = width

        wb.save(output)
        output.seek(0)

        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name='Consultancy_Agreement_List_All.xlsx'
        )

    except Exception as e:
        logger.error(f"Error exporting all contracts to Excel: {str(e)}")
        flash("An error occurred while exporting all contracts to Excel.", 'danger')
        return redirect(url_for('contracts.index'))
@contracts_bp.route('/view/<contract_id>')
@login_required
def view(contract_id):
    try:
        contract = Contract.query.get_or_404(contract_id)
        # Allow admins to view any contract, non-admins only their own
        if not current_user.has_role('admin') and contract.user_id != current_user.id:
            flash("You are not authorized to view this contract.", 'danger')
            return redirect(url_for('contracts.index'))
        if contract.deleted_at is not None:
            flash("This contract has been deleted and cannot be viewed.", 'danger')
            return redirect(url_for('contracts.index'))

        contract_data = contract.to_dict()

        # Format dates
        contract_data['agreement_start_date_display'] = format_date(contract_data['agreement_start_date'])
        contract_data['agreement_end_date_display'] = format_date(contract_data['agreement_end_date'])

        # Get financial data as floats
        total_fee_usd = float(contract_data['total_fee_usd']) if contract_data['total_fee_usd'] else 0.0
        tax_percentage = float(contract_data.get('tax_percentage', 15.0))
        deduct_tax_code = contract_data.get('deduct_tax_code', '')
        vat_organization_name = contract_data.get('vat_organization_name', '')  # New field
        contract_data['total_fee_usd'] = total_fee_usd
        contract_data['gross_amount_usd'] = total_fee_usd
        contract_data['total_fee_words'] = contract_data.get('total_fee_words') or number_to_words(total_fee_usd)

        # Calculate total gross and net as floats
        total_gross_amount, total_net_amount = calculate_payments(
            total_fee_usd, tax_percentage, contract_data.get('payment_installments', [])
        )
        contract_data['total_gross_amount'] = total_gross_amount
        contract_data['total_net_amount'] = total_net_amount
        contract_data['total_gross'] = f"USD{total_gross_amount:.2f}"
        contract_data['total_net'] = f"USD{total_net_amount:.2f}"

        # Process payment installments
        for installment in contract_data.get('payment_installments', []):
            installment['dueDate_display'] = format_date(installment.get('dueDate', ''))
            match = re.search(r'\((\d+\.?\d*)\%\)', installment['description'])
            percentage = float(match.group(1)) if match else 0.0
            gross, tax, net = calculate_installment_payments(total_fee_usd, tax_percentage, percentage)
            installment['gross_amount'] = gross
            installment['tax_amount'] = tax
            installment['net_amount'] = net

        # Define standard articles
        standard_articles = [
            {
                'number': 1,
                'title': 'TERMS OF REFERENCE',
                'content': (
                    '“Party B” shall perform tasks as stated in the attached TOR <strong> (annex-1)</strong> to “Party A”, '
                    'and deliver each milestone as stipulated in article 4.\n\n'
                    'The work shall be of good quality and well performed with the acceptance by “Party A.”'
                ),
                'table': None
            },
            {
                'number': 2,
                'title': 'TERM OF AGREEMENT',
                'content': (
                    f'The agreement is effective from {contract_data["agreement_start_date_display"]} – '
                    f'{contract_data["agreement_end_date_display"]}. This Agreement is terminated automatically '
                    'after the due date of the Agreement Term unless otherwise, both Parties agree to extend '
                    'the Term with a written agreement.'
                ),
                'table': None
            },
            {
                'number': 3,
                'title': 'PROFESSIONAL FEE',
                'content': (
                    f'The professional fee is the total amount of <span style="font-size: 16px;"> <strong> {contract_data["total_gross"]} </strong></span>'
                    f'<span style="font-size: 16px; "><strong> ({contract_data["total_fee_words"]})</strong></span> '
                    f'{"excluding" if tax_percentage == 0 else "including"} tax for the whole assignment period.'
                    f'{"\n\n<span style=\"font-size: 16px; margin-left:40px;\"><strong>" + vat_organization_name + "</strong></span>\n<span style=\"font-size: 16px; margin-left:40px;\"><strong>VAT TIN: " + deduct_tax_code + "</strong></span>" if tax_percentage == 0 and deduct_tax_code and vat_organization_name else ""}\n\n'
                    f'<span style="font-size: 16px; margin-left:40px;"><strong>Total Service Fee: {contract_data["total_gross"]}</strong></span>\n'
                    f'{"<span style=\"font-size: 16px; margin-left:40px;\"><strong>Withholding Tax " + f"{tax_percentage}%: USD{contract_data['total_gross_amount'] * (tax_percentage/100):.2f}</strong></span>\n" if tax_percentage > 0 else ""}'
                    f'<span style="font-size: 16px; margin-left:40px;"><strong>Net amount: {contract_data["total_net"]}</strong></span>\n\n'
                    f'“Party B” is responsible to issue the Invoice (net amount) and receipt (when receiving the payment) '
                    f'with the total amount as stipulated in each instalment as in <strong>Article 4.</strong>\n\n'
                    f'“Party B” is responsible for all related taxes payable to the government department.'
                ),
                'table': None
            },
            {
                'number': 4,
                'title': 'TERM OF PAYMENT',
                'content': 'The payment will be made based on the following schedules:',
                'table': [
                    {'Installment': 'Installment', 'Total Amount (USD)': 'Total Amount (USD)', 'Deliverable': 'Deliverable', 'Due date': 'Due date'},
                    *[
                        {
                            'Installment': installment['description'],
                            'Total Amount (USD)': (
                                f'- Gross: ${installment["gross_amount"]:.2f}\n'
                                f'{"- Tax " + f"{tax_percentage}%: ${installment['tax_amount']:.2f}\n" if tax_percentage > 0 else ""}'
                                f'- Net pay: ${installment["net_amount"]:.2f}'
                            ),
                            'Deliverable': installment['deliverables'].replace('; ', '\n- '),
                            'Due date': installment['dueDate_display']
                        }
                        for installment in contract_data.get('payment_installments', [])
                    ]
                ]
            },
            {
                'number': 5,
                'title': 'NO OTHER PERSONS',
                'content': (
                    'No person or entity, which is not a party to this agreement, has any rights to enforce, '
                    'take any action, or claim it is owed any benefit under this agreement.'
                ),
                'table': None
            },
            {
                'number': 6,
                'title': 'MONITORING and COORDINATION',
                'content': (
                    f'“Party A” shall monitor and evaluate the progress of the agreement toward its objective, '
                    f'including the activities implemented. '
                    f'{" and ".join([f"<strong>{person['name']}</strong>, <strong>{person['position']}</strong> "
                    f"(Telephone {person['phone']} Email: <span style='color: blue; text-decoration: underline;'>{person['email']}</span>)" 
                    for person in contract_data.get("focal_person_info", [])]) or "<strong>N/A</strong>, <strong>N/A</strong> (Telephone N/A Email: N/A)"} '
                    f'is the focal contact person of “Party A” and '
                    f'<strong>{contract_data.get("party_b_signature_name", "N/A")}</strong>, <strong>{contract_data.get("party_b_position", "Freelance Consultant")}</strong> '
                    f'(HP. {contract_data.get("party_b_phone", "N/A")}, E-mail: <span style="color: blue; text-decoration: underline;">{contract_data.get("party_b_email", "N/A")}</span>) '
                    f'the focal contact person of the “Party B”. The focal contact person of “Party A” and “Party B” will work together '
                    f'for overall coordination including reviewing and meeting discussions during the assignment process.'
                ),
                'table': None
            },
            {
                'number': 7,
                'title': 'CONFIDENTIALITY',
                'content': (
                    f'All outputs produced, with the exception of the <strong> “{contract_data.get("project_title", "N/A")}” </strong>, '
                    f'which is a contribution from, and to be claimed as a public document by the main author and co-author '
                    f'in associated, and/or under this agreement, shall be the property of “Party A”. The “Party B” agrees '
                    f'to not disclose any confidential information, of which he/she may take cognizance in the performance '
                    f'under this contract, except with the prior written approval of “Party A”.'
                ),
                'table': None
            },
            {
                'number': 8,
                'title': 'ANTI-CORRUPTION and CONFLICT OF INTEREST',
                'content': (
                    '“Party B” shall not participate in any practice that is or could be construed as an illegal or corrupt '
                    'practice in Cambodia. The “Party A” is committed to fighting all types of corruption and expects this same '
                    'commitment from the consultant it reserves the rights and believes based on the declaration of “Party B” '
                    'that it is an independent social enterprise firm operating in Cambodia and it does not involve any conflict '
                    'of interest with other parties that may be affected to the “Party A”.'
                ),
                'table': None
            },
            {
                'number': 9,
                'title': 'OBLIGATION TO COMPLY WITH THE NGOF’S POLICIES AND CODE OF CONDUCT',
                'content': (
                    'By signing this agreement, “Party B” is obligated to comply with and respect all existing policies and code '
                    'of conduct of “Party A”, such as Gender Mainstreaming, Child Protection, Disability policy, Environmental '
                    'Mainstreaming, etc. and the “Party B” declared themselves that s/he will perform the assignment in the neutral '
                    'position, professional manner, and not be involved in any political affiliation.'
                ),
                'table': None
            },
            {
                'number': 10,
                'title': 'ANTI-TERRORISM FINANCING AND FINANCIAL CRIME',
                'content': (
                    'NGOF is determined that all its funds and resources should only be used to further its mission and shall not '
                    'be subject to illicit use by any third party nor used or abused for any illicit purpose. In order to achieve '
                    'this objective, NGOF will not knowingly or recklessly provide funds, economic goods, or material support to any '
                    'entity or individual designated as a “terrorist” by the international community or affiliate domestic governments '
                    'and will take all reasonable steps to safeguard and protect its assets from such illicit use and to comply with '
                    'host government laws.\nNGOF respects its contracts with its donors and puts procedures in place for compliance '
                    'with these contracts.\n“Illicit use” refers to terrorist financing, sanctions, money laundering, and export '
                    'control regulations.'
                ),
                'table': None
            },
            {
                'number': 11,
                'title': 'INSURANCE',
                'content': (
                    '“Party B” is responsible for any health and life insurance of its team members. “Party A” will not be held '
                    'responsible for any medical expenses or compensation incurred during or after this contract.'
                ),
                'table': None
            },
            {
                'number': 12,
                'title': 'ASSIGNMENT',
                'content': (
                    '“Party B” shall have the right to assign individuals within its organization to carry out the tasks herein '
                    'named in the attached Technical Proposal. The “Party B” shall not assign, or transfer any of its rights or '
                    'obligations under this agreement hereunder without the prior written consent of “Party A”. Any attempt by '
                    '“Party B” to assign or transfer any of its rights and obligations without the prior written consent of “Party A” '
                    'shall render this agreement subject to immediate termination by “Party A”.'
                ),
                'table': None
            },
            {
                'number': 13,
                'title': 'RESOLUTION OF CONFLICTS/DISPUTES',
                'content': (
                    'Conflicts between any of these agreements shall be resolved by the following methods:\n'
                    'In the case of a disagreement arising between “Party A” and the “Party B” regarding the implementation of '
                    'any part of, or any other substantive question arising under or relating to this agreement, the parties shall '
                    'use their best efforts to arrive at an agreeable resolution by mutual consultation.\n'
                    'Unresolved issues may, upon the option of either party and written notice to the other party, be referred to '
                    'for arbitration. Failure by the “Party B” or “Party A” to dispute a decision arising from such arbitration in '
                    'writing within thirty (30) calendar days of receipt of a final decision shall result in such final decision '
                    'being deemed binding upon either the “Party B” and/or “Party A”.<strong> All expenses related to arbitration will be </strong> '
                    '<strong>shared equally between both parties.</strong>'
                ),
                'table': None
            },
            {
                'number': 14,
                'title': 'TERMINATION',
                'content': (
                    'The “Party A” or the “Party B” may, by notice in writing, terminate this agreement under the following conditions:\n\n'
                    '1. “Party A” may terminate this agreement at any time with a one-week notice if “Party B” fails to comply with the '
                    'terms and conditions of this agreement.\n\n'
                    '2. For gross professional misconduct (as defined in the NGOF Human Resource Policy), “Party A” may terminate '
                    'this agreement immediately without prior notice. “Party A” will notify “Party B” in a letter that will indicate '
                    'the reason for termination as well as the effective date of termination.\n\n'
                    '3. “Party B” may terminate this agreement at any time with a one-week notice if “Party A” fails to comply with '
                    'the terms and conditions of this agreement. “Party B” will notify “Party A” in a letter that will indicate the '
                    'reason for termination as well as the effective date of termination. If “Party B” terminates this '
                    'agreement without any appropriate reason or fails to implement the assignment, “Party B” must '
                    'refund the full amount of fees received to “Party A”.\n\n'
                    '4. If for any reason either “Party A” or “Party B” decides to terminate this agreement, “Party B” shall be '
                    'paid pro-rata for the work already completed by “Party A”. This payment will require the submission of a timesheet '
                    'that demonstrates work completed as well as the handing over of any deliverables completed or partially completed. '
                    'In case “Party B” has received payment for services under the agreement which have not yet been performed, the '
                    'appropriate portion of these fees must be refunded by “Party B” to “Party A”.'
                ),
                'table': None
            },
            {
                'number': 15,
                'title': 'MODIFICATION OR AMENDMENT',
                'content': (
                    'No modification or amendment of this agreement shall be valid unless in writing and signed by an authorized '
                    'person of “Party A” and “Party B”.'
                ),
                'table': None
            },
            {
                'number': 16,
                'title': 'CONTROLLING OF LAW',
                'content': (
                    'This agreement shall be governed and construed following the law of the Kingdom of Cambodia. '
                    'The Simultaneous Interpretation Agreement is prepared in two original copies.'
                ),
                'table': None
            }
        ]

        # Prepare custom articles
        custom_articles = [
            {'article_number': str(k), 'custom_sentence': v}
            for k, v in contract_data.get('custom_article_sentences', {}).items()
        ]

        return render_template(
            'contracts/view.html',
            contract=contract_data,
            standard_articles=standard_articles,
            articles=custom_articles,
            format_date=format_date
        )
    except Exception as e:
        logger.error(f"Error viewing contract {contract_id}: {str(e)}")
        flash("An error occurred while viewing the contract.", 'danger')
        return redirect(url_for('contracts.index'))